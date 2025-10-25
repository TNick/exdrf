import logging
from typing import Optional

import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import coordinate_to_tuple, get_column_letter
from openpyxl.worksheet.worksheet import Worksheet
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, Table, TableStyle
from siscadro_util.settings import FontFiles

logger = logging.getLogger(__name__)


class XlWriterBase:
    def __init__(
        self,
        out_file: str = None,
        empty_cell_character: Optional[str] = "",
    ):

        self.stg = None  # stg
        self.out_file = out_file
        self.empty_cell_character = empty_cell_character
        self.wb = None
        self.sht = None
        self.r = -1
        self.htrick_counter = 1
        self.column_widths = []

        self.grey_fill = None
        self.side_thick = None
        self.side_thin = None
        self.font_12_b = None
        self.font_12 = None
        self.font_11 = None
        self.font_11_b = None
        self.font_10 = None
        self.font_10_b = None
        self.font_09 = None
        self.font_08 = None

        self.border_thin = None
        self.border_thick = None
        self.align_center = None
        self.align_left = None
        self.align_right = None

    def cell(self, name, r=None, c=None):
        r = r if r > 0 else self.r - r
        return self.sht.cell(row=r, column=c)

    def set_cell(
        self,
        name=None,
        r=None,
        c=None,
        mr=1,
        mc=1,
        value=None,
        style=None,
        htrick=False,
        empty_character=None,
    ):
        r = r if r > 0 else self.r - r
        # print(f"set_cell: {name}, r={r}, c={c}, mr={mr}, mc={mc}, value={value}, style={style}, htrick={htrick}")
        cell = self.sht.cell(row=r, column=c)
        if value is None or value == "":
            if empty_character is None:
                value = self.empty_cell_character
            else:
                value = empty_character
        cell.value = value
        if mr > 1 or mc > 1:
            self.sht.merge_cells(
                start_row=r,
                start_column=c,
                end_row=r + mr - 1,
                end_column=c + mc - 1,
            )
            if style:
                for rs in range(0, mr):
                    for cs in range(0, mc):
                        self.sht.cell(row=r + rs, column=c + cs).style = (
                            style if isinstance(style, str) else style.name
                        )
        elif style:
            cell.style = style if isinstance(style, str) else style.name
        if htrick and cell.value:
            self.htrick_counter = self.htrick_counter + 1
            ghost = self.sht.cell(
                row=cell.row,
                column=cell.column
                + len(self.column_widths)
                + self.htrick_counter,
            )
            # ghost.width = sum(self.column_widths)
            cd = self.sht.column_dimensions[get_column_letter(ghost.column)]
            cd.width = self.merged_width(c, mc)
            cd.auto_size = False
            cd.hidden = True
            if style:
                ghost.style = style if isinstance(style, str) else style.name
            ghost.value = cell.value
        return cell

    def merged_width(self, col, count):
        return sum(self.column_widths[col - 1 : col + count - 1])

    def set_ts_cell(
        self,
        name,
        r=None,
        c=None,
        mr=1,
        mc=1,
        default=None,
        args=None,
        kwargs=None,
        style=None,
        htrick=False,
    ):
        return self.set_cell(
            name,
            r=r,
            c=c,
            mr=mr,
            mc=mc,
            # value=self.ts(name, default, *args if args else [], **kwargs if kwargs else {}),
            value=default.format(
                *args if args else [], **kwargs if kwargs else {}
            ),
            style=style,
            htrick=htrick,
        )

    def prepare(self):
        self.sht = self.wb.active
        self.r = 1
        # self.sht.sheet_properties.pageSetUpPr.fitToPage = True
        # self.sht.page_setup.fitToPage = True
        self.sht.print_options.horizontalCentered = True
        self.sht.print_options.verticalCentered = False
        self.sht.page_setup.orientation = self.sht.ORIENTATION_LANDSCAPE
        self.sht.page_setup.paperSize = self.sht.PAPERSIZE_A4
        self.sht.page_setup.fitToWidth = True
        self.sht.page_setup.fitToHeight = False
        self.sht.page_margins.left = 0.4
        self.sht.page_margins.right = 0.2
        self.sht.page_margins.top = 0.2
        self.sht.page_margins.bottom = 0.2
        self.sht.page_margins.header = 0.0
        self.sht.page_margins.footer = 0.0

    def predefined_base(self):

        self.grey_fill = PatternFill(
            fill_type="solid", start_color="FFEEEEEE", end_color="FFEEEEEE"
        )
        self.side_thick = Side(style="thick", color="000000")
        self.side_thin = Side(style="thin", color="000000")
        self.font_12_b = Font(bold=True, size=12)
        self.font_12 = Font(bold=False, size=12)
        self.font_11 = Font(bold=False, size=11)
        self.font_11_b = Font(bold=True, size=11)
        self.font_10 = Font(bold=False, size=10)
        self.font_10_b = Font(bold=True, size=10)
        self.font_09 = Font(bold=False, size=9)
        self.font_08 = Font(bold=False, size=8)

        self.border_thin = Border(
            left=self.side_thin,
            top=self.side_thin,
            right=self.side_thin,
            bottom=self.side_thin,
            vertical=self.side_thin,
            horizontal=self.side_thin,
        )
        self.border_thick = Border(
            left=self.side_thick,
            top=self.side_thick,
            right=self.side_thick,
            bottom=self.side_thick,
            vertical=self.side_thick,
            horizontal=self.side_thick,
        )

        self.align_center = Alignment(
            horizontal="center",
            vertical="center",
            text_rotation=0,
            wrap_text=True,
            shrink_to_fit=False,
            indent=0,
        )
        self.align_left = Alignment(
            horizontal="left",
            vertical="center",
            text_rotation=0,
            wrap_text=True,
            shrink_to_fit=False,
            indent=0,
        )
        self.align_right = Alignment(
            horizontal="right",
            vertical="center",
            text_rotation=0,
            wrap_text=True,
            shrink_to_fit=False,
            indent=0,
        )

    def predefined_styles(self):
        pass

    def generate(self):
        """
        Generates the Excel file.
        """
        raise NotImplementedError

    def write(self):
        """
        Writes the content of the workbook to the file.
        """
        if self.out_file:
            self.wb.save(self.out_file)
        else:
            raise ValueError("No file path was set")
        logger.debug("xls output written to %s", self.out_file)

    def prepare_pdf_for_table(self):
        par_sty_center = ParagraphStyle("Normal", None, alignment=TA_CENTER)
        par_sty_left = ParagraphStyle("Normal", None, alignment=TA_LEFT)
        par_sty_right = ParagraphStyle("Normal", None, alignment=TA_RIGHT)

        styles = [
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            # ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
            # ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
            ("FONTNAME", (0, 0), (-1, -1), "DejaVuSerifCondensed"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
        ]
        return par_sty_center, par_sty_left, par_sty_right, styles

    def range_to_pdf_table(self, sht, sr, sc, er, ec, default_width=50):
        """
        Creates a table in a pdf document from excel data.

        :param sht: source sheet;
        :param sr: start row in Excel (1 based);
        :param sc: start column in excel (1 based);
        :param er: end row in Excel (1 based);
        :param ec: end column in Excel (1 based);
        :param default_width: default width of a column;
        :return: the table that was created
        """
        par_sty_center, par_sty_left, par_sty_right, styles = (
            self.prepare_pdf_for_table()
        )
        data = []
        images = {}
        for image in sht._images:
            r, c = coordinate_to_tuple(image.anchor)
            if sr <= r <= er and sc <= c <= ec:
                images[f"{r}-{c}"] = (
                    image.full_path,
                    image.width * 0.8,
                    image.height * 0.8,
                )

        for r in range(sr, er + 1):
            row = []
            for c in range(sc, ec + 1):
                image_key = f"{r}-{c}"
                par_sty = par_sty_center

                cell = sht.cell(r, c)
                value = cell.value if cell.value is not None else ""
                if (
                    cell.border.left and cell.border.left.style is not None
                ) or (cell.border.top and cell.border.top.style is not None):
                    styles.append(
                        (
                            "INNERGRID",
                            (c - sc, r - sr),
                            (c - sc, r - sr),
                            0.25,
                            colors.black,
                        )
                    )
                    styles.append(
                        (
                            "BOX",
                            (c - sc, r - sr),
                            (c - sc, r - sr),
                            0.25,
                            colors.black,
                        )
                    )
                if image_key in images:
                    image = images[image_key]
                    value = (
                        f'{"" if value is None else str(value)}<img '
                        f'src="{image[0]}" '
                        f'width="{image[1]}" '
                        f'height="{image[2]}" valign="middle"/>'
                    )
                if len(str(value)):
                    value = str(value)
                    if len(row) == 0 and r == 53:
                        if "Din care:" == value:
                            pass
                    if 0 and cell.font.name:
                        value = f'<font name="{cell.font.name}" size="{cell.font.sz * 0.8}">{value}</font>'
                    else:
                        value = f'<font name="DejaVuSerifCondensed" size="{cell.font.sz * 0.8}">{value}</font>'
                    # if cell.font.sz != 10.0:
                    #     styles.append(('FONTSIZE', (c-sc, r-sr), (c-sc, r-sr), cell.font.sz*0.9))
                    #     value = f'<font size="{cell.font.sz*0.9}">{value}</font>'
                    if cell.font.color:
                        try:
                            color = "0x{}".format(cell.font.color[2:])
                        except TypeError:
                            color = "0x000000"
                        styles.append(
                            (
                                "TEXTCOLOR",
                                (c - sc, r - sr),
                                (c - sc, r - sr),
                                color,
                            )
                        )
                    if cell.font.bold:
                        value = "<b>{}</b>".format(value)
                    if cell.font.italic:
                        value = "<i>{}</i>".format(value)
                    if cell.font.underline:
                        value = "<u>{}</u>".format(value)
                    if cell.fill.bgColor:
                        if cell.fill.bgColor.rgb != "00000000":
                            styles.append(
                                (
                                    "BACKGROUND",
                                    (c - sc, r - sr),
                                    (c - sc, r - sr),
                                    "0x{}".format(cell.fill.bgColor.rgb[2:]),
                                )
                            )

                    if cell.alignment.horizontal:
                        if cell.alignment.horizontal == "left":
                            par_sty = par_sty_left
                        elif cell.alignment.horizontal == "right":
                            par_sty = par_sty_right
                        styles.append(
                            (
                                "ALIGN",
                                (c - sc, r - sr),
                                (c - sc, r - sr),
                                cell.alignment.horizontal.upper(),
                            )
                        )
                    if cell.alignment.vertical:
                        styles.append(
                            (
                                "VALIGN",
                                (c - sc, r - sr),
                                (c - sc, r - sr),
                                cell.alignment.vertical.upper(),
                            )
                        )
                    # if cell.style:
                    #     style = self.sht.parent._named_styles[cell.style]
                    #     font = style.font
                    #     alignment = style.alignment
                    #     border = style.border
                    #     fill = style.fill
                    # else:
                    #     pass
                    row.append(Paragraph(value, par_sty))
                else:
                    row.append("")
            data.append(row)

        col_w = []
        for c in range(sc, ec + 1):
            cw = self.sht.column_dimensions[get_column_letter(c)].width * 5.2
            if not cw:
                cw = default_width
            col_w.append(cw)
        for merged in sht.merged_cells:
            if merged.min_row > er or merged.max_row > er:
                continue
            if merged.min_row < sr or merged.max_row < sr:
                continue
            if merged.min_col > ec or merged.max_col > ec:
                continue
            if merged.min_col < sc or merged.max_col < sc:
                continue
            value = []
            for r in range(merged.min_row, merged.max_row + 1):
                for c in range(merged.min_col, merged.max_col + 1):
                    cell = data[r - sr][c - sc]
                    if cell != "":
                        value.append(cell)
                        data[r - sr][c - sc] = ""

            if len(value) == 0:
                value = ""
            elif len(value) == 1:
                value = value[0]
            else:
                value = value
            # value = [Paragraph('+', par_sty), value]
            # for r in range(merged.min_row, merged.max_row + 1):
            #     for c in range(merged.min_col, merged.max_col + 1):
            #         data[r-sr][c-sc] = value

            data[merged.min_row - sr][merged.min_col - sc] = value
            styles.append(
                (
                    "SPAN",
                    (merged.min_col - sc, merged.min_row - sr),
                    (merged.max_col - sc, merged.max_row - sr),
                ),
            )

        table = Table(data, colWidths=col_w)
        table.setStyle(TableStyle(styles))

        return table

    def save_as_pdf(self, path):
        """
        Export the file to pdf.
        """
        raise NotImplementedError

    def prepare_pdf(self):
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        from reportlab.pdfbase.ttfonts import TTFont

        # from reportlab.lib import colors

        def do_one_font(name, data: FontFiles):
            pdfmetrics.registerFont(TTFont(name, data.regular, "UTF-8"))
            pdfmetrics.registerFont(TTFont(f"{name}-Bold", data.bold, "UTF-8"))
            pdfmetrics.registerFont(
                TTFont(f"{name}-Italic", data.italic, "UTF-8")
            )
            pdfmetrics.registerFont(
                TTFont(f"{name}-BoldItalic", data.bold_italic, "UTF-8")
            )
            registerFontFamily(
                name,
                normal=name,
                bold=f"{name}-Bold",
                italic=f"{name}-Italic",
                boldItalic=f"{name}-BoldItalic",
            )

        do_one_font("DejaVuSerif", self.stg.fonts.dejavu_serif)
        do_one_font(
            "DejaVuSerifCondensed", self.stg.fonts.dejavu_serif_condensed
        )

    def fill_empty(self, r, c, rc, cc, style=None):
        """Fills with the empty character a range of cells."""
        if style is None:
            style = self.style_table_cell_center
        for ir in range(rc):
            for ic in range(cc):
                self.set_cell(
                    r=r + ir,
                    c=c + ic,
                    value=self.empty_cell_character,
                    style=style,
                )

    def save_as_doc(self, doc: Optional[docx.Document] = None):
        """Saves the table into a word document."""
        if doc is None:
            doc = docx.Document()

        raise NotImplementedError

    def range_to_docx_table(
        self,
        sht: Worksheet,
        doc: docx.Document,
        sr,
        sc,
        er,
        ec,
        default_width=50,
        style="Table Grid",
        para_style=None,
    ):
        """
        Creates a table in a pdf document from excel data.

        :param sht: The source;
        :param doc: The document where we add the table;
        :param sr: start row in Excel (1 based);
        :param sc: start column in excel (1 based);
        :param er: end row in Excel (1 based);
        :param ec: end column in excel (1 based);
        :return: the table that was created
        """
        table = doc.add_table(rows=er - sr + 1, cols=ec - sc + 1, style=style)

        table_cells = table._cells
        # for c in range(sc, ec + 1):
        #     cw = self.sht.column_dimensions[get_column_letter(c)].width * 5.2
        #     if not cw:
        #         cw = default_width

        for merged in sht.merged_cells:
            if merged.min_row > er or merged.max_row > er:
                continue
            if merged.min_row < sr or merged.max_row < sr:
                continue
            if merged.min_col > ec or merged.max_col > ec:
                continue
            if merged.min_col < sc or merged.max_col < sc:
                continue

            top_left = table_cells[
                (merged.min_row - sr) * (ec - sc + 1) + (merged.min_col - sc)
            ]
            bottom_right = table_cells[
                (merged.max_row - sr) * (ec - sc + 1) + (merged.max_col - sc)
            ]
            merged = top_left.merge(bottom_right)

        for r in range(sr, er + 1):
            for c in range(sc, ec + 1):
                cell = sht.cell(r, c)
                value = cell.value if cell.value is not None else ""
                # dst = table.rows[r-sr].cells[c-ec]
                dst = table_cells[(r - sr) * (ec - sc + 1) + (c - sc)]
                dst.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if len(str(value)):
                    # dst.text = str(value)
                    paragraph = dst.paragraphs[0]
                    if para_style:
                        paragraph.style = doc.styles[para_style]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(str(value))
                    run.bold = cell.font.bold
                    run.italic = cell.font.italic
                    run.underline = cell.font.underline
                    if cell.alignment.horizontal:
                        if cell.alignment.horizontal == "left":
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif cell.alignment.horizontal == "right":
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if cell.alignment.vertical:
                        if cell.alignment.vertical == "top":
                            paragraph.alignment = WD_ALIGN_VERTICAL.TOP
                        elif cell.alignment.vertical == "bottom":
                            paragraph.alignment = WD_ALIGN_VERTICAL.BOTTOM
