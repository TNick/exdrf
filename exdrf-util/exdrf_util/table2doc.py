import logging
from typing import TYPE_CHECKING

from attrs import define, field

from exdrf_util.table2base import Range2Other

if TYPE_CHECKING:
    from docx.document import Document
    from docx.table import Table as DocxTable
    from openpyxl.worksheet.cell_range import CellRange  # type: ignore

logger = logging.getLogger(__name__)


@define
class Range2Docx(Range2Other):

    document: "Document"
    table_style: str = field(default="Table Grid")
    para_style: str = field(default="Normal")
    table: "DocxTable" = field(default=None, init=False)

    def generate(self) -> "DocxTable":
        """Generate the pdf table."""

        self.create_layout()

        for self.crt_row in range(self.start_row, self.last_row + 1):
            for self.crt_col in range(self.start_col, self.last_col + 1):
                self.create_cell()

        return self.table

    def create_layout(self):
        # Create the table in the document with all cells.
        self.table = self.document.add_table(
            self.row_count, self.col_count, style=self.table_style
        )

        # Center the table and make it span the full available page width.
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Emu

        section = self.document.sections[-1]
        page_width_len = section.page_width
        left_margin_len = section.left_margin
        right_margin_len = section.right_margin

        page_width_emu = (
            int(page_width_len) if page_width_len is not None else 0
        )
        left_margin_emu = (
            int(left_margin_len) if left_margin_len is not None else 0
        )
        right_margin_emu = (
            int(right_margin_len) if right_margin_len is not None else 0
        )
        available_width_emu = (
            page_width_emu - left_margin_emu - right_margin_emu
        )

        self.table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.table.autofit = False

        if self.col_count > 0:
            col_width_emu = max(1, available_width_emu // self.col_count)
            for column in self.table.columns:
                column.width = Emu(col_width_emu)

        for merged in self.worksheet.merged_cells:
            self.merge_cells(merged)

    def merge_cells(self, merged: "CellRange"):
        """Merge the cells in the table."""
        if merged.min_row > self.last_row or merged.max_row > self.last_row:
            return
        if merged.min_row < self.start_row or merged.max_row < self.start_row:
            return
        if merged.min_col > self.last_col or merged.max_col > self.last_col:
            return
        if merged.min_col < self.start_col or merged.max_col < self.start_col:
            return

        table_cells = self.table._cells

        tl_index = (merged.min_row - self.start_row) * self.col_count + (
            merged.min_col - self.start_col
        )

        br_index = (merged.max_row - self.start_row) * self.col_count + (
            merged.max_col - self.start_col
        )

        if tl_index >= len(table_cells) or br_index >= len(table_cells):
            return

        try:
            top_left = table_cells[tl_index]
            bottom_right = table_cells[br_index]
            top_left.merge(bottom_right)
        except IndexError:
            logger.warning("Index error merging cells: %s", merged)
            return
        except Exception:
            logger.exception("Exception merging cells: %s", merged)
            return

    @property
    def doc_cell(self):
        """Get the current cell from the document."""
        index = self.rel_row * self.col_count + self.rel_col
        try:
            return self.table._cells[index]
        except IndexError:
            logger.warning("Index error getting cell: %s", index)
            return None

    def create_cell(self):
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        xl_cell = self.xl_cell
        doc_cell = self.doc_cell
        value = xl_cell.value if xl_cell.value is not None else ""
        str_value = str(value)

        if not str_value:
            return

        paragraph = doc_cell.paragraphs[0]  # type: ignore
        if self.para_style:
            paragraph.style = self.para_style

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(value))
        run.bold = xl_cell.font.bold  # type: ignore
        run.italic = xl_cell.font.italic  # type: ignore
        run.underline = xl_cell.font.underline  # type: ignore
        if xl_cell.alignment.horizontal:  # type: ignore
            if xl_cell.alignment.horizontal == "left":  # type: ignore
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif xl_cell.alignment.horizontal == "right":  # type: ignore
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if xl_cell.alignment.vertical:  # type: ignore
            if xl_cell.alignment.vertical == "top":  # type: ignore
                paragraph.alignment = WD_ALIGN_VERTICAL.TOP  # type: ignore
            elif xl_cell.alignment.vertical == "bottom":  # type: ignore
                paragraph.alignment = WD_ALIGN_VERTICAL.BOTTOM  # type: ignore

    def remove_table_borders(self):
        # Remove borders from all cells
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        for row in self.table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for border_name in (
                    "top",
                    "left",
                    "bottom",
                    "right",
                    "insideH",
                    "insideV",
                ):
                    border_el = OxmlElement(f"w:{border_name}")
                    border_el.set(qn("w:val"), "nil")
                    tcBorders.append(border_el)
                tcPr.append(tcBorders)
