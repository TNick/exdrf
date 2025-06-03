import os
import re
import tempfile

import cairosvg
from bs4 import BeautifulSoup, Tag
from bs4.element import NavigableString
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


def svg_to_png(svg_str):
    tmp_fd, tmp_png_path = tempfile.mkstemp(suffix=".png")
    os.close(tmp_fd)
    cairosvg.svg2png(bytestring=svg_str.encode("utf-8"), write_to=tmp_png_path)
    return tmp_png_path


def parse_color(color_str):
    # Supports rgb(255,0,0), #RRGGBB, and named colors (limited)
    if not color_str:
        return None
    color_str = color_str.strip()
    if color_str.startswith("#") and len(color_str) == 7:
        r = int(color_str[1:3], 16)
        g = int(color_str[3:5], 16)
        b = int(color_str[5:7], 16)
        return RGBColor(r, g, b)
    elif color_str.startswith("rgb"):
        nums = [int(n) for n in re.findall(r"\d+", color_str)]
        if len(nums) == 3:
            return RGBColor(*nums)
    # Limited named colors
    COLORS = {
        "red": RGBColor(255, 0, 0),
        "black": RGBColor(0, 0, 0),
        "white": RGBColor(255, 255, 255),
        "yellow": RGBColor(255, 255, 0),
        "green": RGBColor(0, 128, 0),
        "blue": RGBColor(0, 0, 255),
    }
    return COLORS.get(color_str.lower(), None)


def set_cell_shading(cell, color):
    if not color:
        return
    color_hex = f"{color.rgb:06x}"
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {qn("w:fill")}="{color_hex}"/>')
    )


def set_cell_border(cell, color):
    if not color:
        return
    color_hex = f"{color.rgb:06x}"
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_xml = (
        f"<w:tcBorders>"
        f'  <w:top w:val="single" w:sz="4" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="4" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="{color_hex}"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(parse_xml(borders_xml))


def handle_inline_formatting(p, node):
    """Recursively add runs to p for inline formatting in a node."""

    def apply_format(run, tag):
        if tag.name in ["b", "strong"]:
            run.bold = True
        if tag.name in ["i", "em"]:
            run.italic = True
        if tag.name in ["u"]:
            run.underline = True
        if tag.name in ["s", "strike", "del"]:
            run.font.strike = True

    if isinstance(node, NavigableString):
        if node.strip():
            p.add_run(node)
    elif isinstance(node, Tag):
        run = p.add_run()
        apply_format(run, node)
        # Color and font from style attribute
        style = node.get("style", "")
        color_match = re.search(r"color\s*:\s*([^;]+)", style)
        if color_match:
            color = parse_color(color_match.group(1))
            if color:
                run.font.color.rgb = color
        # Font-weight, etc.
        if "font-weight: bold" in style:
            run.bold = True
        if "font-style: italic" in style:
            run.italic = True
        if "text-decoration: underline" in style:
            run.underline = True
        if "text-decoration: line-through" in style:
            run.font.strike = True
        for child in node.children:
            handle_inline_formatting(p, child)


def handle_table(doc, table_tag):
    # Get max columns, accounting for colspan
    rows = table_tag.find_all("tr")
    max_cols = 0
    for r in rows:
        c = 0
        for td in r.find_all(["td", "th"]):
            colspan = int(td.get("colspan", 1))
            c += colspan
        max_cols = max(max_cols, c)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    # Track which cells are occupied (for rowspan/colspan)
    occupancy = [[None for _ in range(max_cols)] for _ in range(len(rows))]
    for row_idx, row in enumerate(rows):
        col_idx = 0
        for cell in row.find_all(["td", "th"]):
            # Find next available cell
            while col_idx < max_cols and occupancy[row_idx][col_idx]:
                col_idx += 1
            colspan = int(cell.get("colspan", 1))
            rowspan = int(cell.get("rowspan", 1))
            # Insert text with formatting
            docx_cell = table.cell(row_idx, col_idx)
            docx_cell.text = ""
            p = docx_cell.paragraphs[0]
            for child in cell.contents:
                handle_inline_formatting(p, child)
            # Set background color and border
            style = cell.get("style", "")
            bg_match = re.search(r"background(-color)?\s*:\s*([^;]+)", style)
            border_match = re.search(
                r"border.*:\s*[^;]*([#a-zA-Z0-9\(,\)\s]+)", style
            )
            bg_color = parse_color(bg_match.group(2)) if bg_match else None
            border_color = (
                parse_color(border_match.group(1)) if border_match else None
            )
            if bg_color:
                set_cell_shading(docx_cell, bg_color)
            if border_color:
                set_cell_border(docx_cell, border_color)
            # Mark cell as occupied for rowspan/colspan
            for dr in range(rowspan):
                for dc in range(colspan):
                    if row_idx + dr < len(rows) and col_idx + dc < max_cols:
                        occupancy[row_idx + dr][col_idx + dc] = True
            # Merge cells for colspan
            if colspan > 1:
                docx_cell.merge(table.cell(row_idx, col_idx + colspan - 1))
            # Merge cells for rowspan
            if rowspan > 1:
                docx_cell.merge(table.cell(row_idx + rowspan - 1, col_idx))
            col_idx += colspan

    doc.add_paragraph("")


def handle_list(doc, ul_tag, is_ordered=False):
    for li in ul_tag.find_all("li", recursive=False):
        p = doc.add_paragraph(
            style="List Number" if is_ordered else "List Bullet"
        )
        for child in li.children:
            handle_inline_formatting(p, child)


def process_node(doc, node):
    if isinstance(node, NavigableString):
        if node.strip():
            doc.add_paragraph(node)
    elif isinstance(node, Tag):
        if node.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(node.name[1])
            p = doc.add_heading(level=level)
            for child in node.children:
                handle_inline_formatting(p, child)
        elif node.name == "p":
            p = doc.add_paragraph()
            for child in node.children:
                handle_inline_formatting(p, child)
        elif node.name == "ul":
            handle_list(doc, node, is_ordered=False)
        elif node.name == "ol":
            handle_list(doc, node, is_ordered=True)
        elif node.name == "table":
            handle_table(doc, node)
        elif node.name == "img" and node.get("src"):
            src = node["src"]
            if src.startswith("data:image"):
                import base64

                head, data = src.split(",", 1)
                ext = head.split("/")[1].split(";")[0]
                tmp_fd, tmp_img_path = tempfile.mkstemp(suffix=f".{ext}")
                os.close(tmp_fd)
                with open(tmp_img_path, "wb") as f:
                    f.write(base64.b64decode(data))
                doc.add_picture(tmp_img_path, width=Inches(5))
                os.remove(tmp_img_path)
            else:
                if os.path.exists(src):
                    doc.add_picture(src, width=Inches(5))
        elif node.name == "svg":
            svg_str = str(node)
            png_path = svg_to_png(svg_str)
            doc.add_picture(png_path, width=Inches(5))
            os.remove(png_path)
        elif node.name in ["br"]:
            doc.add_paragraph()
        else:
            for child in node.contents:
                process_node(doc, child)


def html_to_docx(html_path, output_docx):
    with open(html_path, encoding="utf-8") as f:
        soup = BeautifulSoup(f, "lxml")
    doc = Document()
    body = soup.body if soup.body else soup
    for node in body.children:
        process_node(doc, node)
    doc.save(output_docx)


if __name__ == "__main__":
    # Usage
    html_to_docx(
        r"D:\prog\CadPlatform\bk-one\playground\yyy.html",
        r"D:\prog\CadPlatform\bk-one\playground\output-yyy.docx",
    )
