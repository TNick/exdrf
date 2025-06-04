import io
import json  # For JS communication
import logging
import re
from copy import deepcopy
from typing import TYPE_CHECKING, Any, List, Union

from attrs import define, field
from bs4 import BeautifulSoup, FeatureNotFound, Tag
from bs4.element import NavigableString, PageElement  # Corrected import

# Imports for HtmlToDocxConverter
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Inches, Pt, RGBColor
from lxml.etree import _Element  # type: ignore
from minify_html import minify
from PyQt5.QtCore import QBuffer, QByteArray, QEvent, QIODevice, Qt, QUrl
from PyQt5.QtGui import QDesktopServices, QImage
from PyQt5.QtWebEngineWidgets import QWebEngineView
from PyQt5.QtWidgets import QApplication  # For processEvents

from exdrf_qt.context_use import QtUseContext

if TYPE_CHECKING:
    from exdrf_qt.context import QtContext  # noqa: F401

OxmlElementType = Union[BaseOxmlElement, _Element]
logger = logging.getLogger(__name__)

# Define constants for styling
# Store colors as (R, G, B) tuples for opacity processing
# And create RGBColor objects on demand.
BORDER_COLOR_SUCCESS_RGB = (0x19, 0x87, 0x54)  # Bootstrap's success green
TABLE_STRIPED_BG_COLOR_RGB = (0xF2, 0xF2, 0xF2)  # Light gray for striping
DEFAULT_BORDER_COLOR_RGB = (0x00, 0x00, 0x00)  # Black
DEFAULT_BORDER_WIDTH_PT = 0.75  # Approx 1px, use in Pt() later


class WebView(QWebEngineView, QtUseContext):
    """A custom web view that can handle internal navigation requests."""

    def __init__(self, ctx: "QtContext", *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.ctx = ctx
        self.devtools_view = None
        self.installEventFilter(self)

    def eventFilter(self, obj, event):  # type: ignore
        """Handle mouse press events."""
        if event.type() == QEvent.Type.MouseButtonPress:
            history = self.history()
            if history is None:
                return super().eventFilter(obj, event)

            # Handle the 'Back' button
            if event.button() == Qt.MouseButton.BackButton:
                if history.canGoBack():
                    history.back()
                event.accept()
                return True  # Return True as event is handled

            # Handle the 'Forward' button
            if event.button() == Qt.MouseButton.ForwardButton:
                if history.canGoForward():
                    history.forward()
                event.accept()
                return True  # Return True as event is handled

        # Otherwise, handle as usual
        return super().eventFilter(obj, event)

    def show_devtools(self):
        if not self.devtools_view:
            # DevTools view
            self.devtools_view = QWebEngineView()
            self.devtools_view.setWindowTitle("DevTools")
            page = self.page()
            assert page is not None
            page.setDevToolsPage(self.devtools_view.page())
            self.devtools_view.closeEvent = (
                lambda event: self._on_devtools_close(event)  # type: ignore
            )
        self.devtools_view.show()
        self.devtools_view.raise_()

    def closeEvent(self, event):  # type: ignore
        if self.devtools_view:
            self.devtools_view.close()
            self.devtools_view = None
        super().closeEvent(event)

    def _on_devtools_close(self, event):  # type: ignore
        self.devtools_view = None
        event.accept()


class HtmlToDocxConverter:
    output_path: str
    view: QWebEngineView
    doc: DocumentObject
    image_cache: dict[str, bytes]
    full_page_qimage: QImage | None
    elements_map: dict[str, dict[str, Any]]
    page_width: int
    page_height: int

    def __init__(self, view: QWebEngineView):  # view is now mandatory
        self.view = view
        self.doc = Document()
        self.image_cache: dict[str, bytes] = {}
        self.full_page_qimage: QImage | None = None
        # Initialize new elements_map
        # element_id: {
        #   "tagName": str, "is_hidden": bool, "geometry": Optional[dict]
        # }
        self.elements_map: dict[str, dict[str, Any]] = {}

    def export_to_docx(self, output_path: str):
        self.output_path = output_path
        logger.debug("Starting DOCX export to: %s", output_path)
        self.prepare_assets()

    def prepare_assets(self):
        logger.debug(
            "Preparing assets (JS execution for element info and screenshot)..."
        )
        # Corrected string literal definition for JS code
        js_get_geometries = """
        (function() {
            console.log(
                "Preparing assets: identifying elements and their visibility"
            );
            let selector = 'p, div, h1, h2, h3, h4, h5, h6, ul, ol, li, dl, ' +
                'dt, dd, table, tr, td, th, thead, tbody, tfoot, span, img, ' +
                'svg, hr, a, b, i, strong, em, u, s, strike, del, font, ' +
                'sup, sub';
            let elements = document.querySelectorAll(selector);
            console.log(
                "Found", elements.length, "relevant elements for d-none check"
            );

            let elementInfoList = [];
            let scrollX = window.scrollX;
            let scrollY = window.scrollY;

            for (let i = 0; i < elements.length; i++) {
                let el = elements[i];
                let id = 'docgen_elem_' + i;
                el.setAttribute('data-docgen-id', id);

                let classList = el.classList;
                let isHiddenByDNone = classList.contains('d-none');
                let hasCollapse = classList.contains('collapse');
                let hasShow = classList.contains('show');

                let isHiddenByCollapseLogic = false;
                if (hasCollapse && !hasShow) {
                    isHiddenByCollapseLogic = true;
                }

                let isEffectivelyHidden =
                    isHiddenByDNone || isHiddenByCollapseLogic;

                let geometry = null;
                let tagName = el.tagName.toLowerCase();

                if (tagName === 'img' || tagName === 'svg') {
                    let rect = el.getBoundingClientRect();
                    if (rect.width > 0 && rect.height > 0) {
                        geometry = {
                            docX: rect.left + scrollX,
                            docY: rect.top + scrollY,
                            docWidth: rect.width,
                            docHeight: rect.height
                        };
                    } else {
                         console.log(
                            "Element " + id + " (" + tagName +
                            ") has zero width/height, no geometry stored."
                        );
                    }
                }
                elementInfoList.push({
                    id: id,
                    tagName: tagName,
                    isEffectivelyHidden: isEffectivelyHidden,
                    geometry: geometry
                });
            }

            let scrollWidth = document.documentElement.scrollWidth;
            let scrollHeight = document.documentElement.scrollHeight;
            console.log("Full page dimensions:", scrollWidth, scrollHeight);

            let final_result = {
                elementInfoList: elementInfoList,
                pageDimensions: { width: scrollWidth, height: scrollHeight }
            };
            console.log(
                "JS result sample:",
                elementInfoList.length > 0 ? elementInfoList[0] : "No elements"
            );
            return JSON.stringify(final_result);
        })();
        """  # Corrected string literal definition for JS code
        logger.debug("Running JS for geometries and visibility...")
        self._run_js_async(js_get_geometries, self.assets_prepared)

    def assets_prepared(self, js_result_str):
        logger.debug("JS execution finished. assets_prepared called.")
        if not js_result_str:
            logger.error("Failed to get element information from JavaScript.")
            return

        try:
            js_data = json.loads(js_result_str)
        except json.JSONDecodeError as e:
            logger.error(
                f"Failed to parse JS result: {e}. Result was: {js_result_str}"
            )
            return
        logger.debug("Successfully parsed JS result.")

        # Process elementInfoList into self.elements_map
        raw_element_info = js_data.get("elementInfoList", [])
        page_dims = js_data.get("pageDimensions")

        self.elements_map = {}
        for item in raw_element_info:
            self.elements_map[item["id"]] = {
                "tagName": item["tagName"],
                "is_hidden": item["isEffectivelyHidden"],
                "geometry": item.get("geometry"),  # Will be None if not present
            }

        # Debug: Log a sample from the processed map
        if self.elements_map:
            sample_id = next(iter(self.elements_map))
            logger.debug(f"Py map {sample_id}: {self.elements_map[sample_id]}")
        logger.debug(
            "Processed %d elements into self.elements_map.",
            len(self.elements_map),
        )

        if not page_dims or not self.view:
            logger.error(
                "Could not get page dimensions or view not available "
                "for screenshot."
            )
            return

        self.page_width = int(page_dims["width"])
        self.page_height = int(page_dims["height"])
        logger.debug(
            "Page dimensions from JS: %dx%s", self.page_width, self.page_height
        )

        # Ensure minimum size for grab
        self.page_width = max(
            self.page_width, self.view.minimumSizeHint().width(), 100
        )
        self.page_height = max(
            self.page_height, self.view.minimumSizeHint().height(), 100
        )
        logger.debug(
            "Adjusted page dimensions for grab: %dx%s",
            self.page_width,
            self.page_height,
        )

        self._run_js_async("window.scrollTo(0,0);", self.window_scrolled)

    def window_scrolled(self, js_result_str):
        logger.debug("Window scrolled to (0,0). Taking screenshot...")
        original_size = self.view.size()
        QApplication.processEvents()

        self.view.resize(self.page_width, self.page_height)
        QApplication.processEvents()  # Allow resize and layout
        # A more robust wait might be needed here, e.g., QTimer

        pixmap = self.view.grab()
        self.full_page_qimage = pixmap.toImage()
        assert self.full_page_qimage is not None

        self.view.resize(original_size)  # Restore original size
        QApplication.processEvents()

        if self.full_page_qimage.isNull():
            logger.error("Failed to grab full page screenshot.")
            self.full_page_qimage = None
        else:
            logger.info(
                "Full page screenshot captured: %sx%s",
                self.full_page_qimage.width(),
                self.full_page_qimage.height(),
            )
        logger.debug("Requesting HTML content with IDs...")
        page = self.view.page()
        assert page is not None
        page.toHtml(self.got_current_html_with_ids)

    def got_current_html_with_ids(self, html_content: str):
        logger.debug(
            "Received HTML content. Starting DOCX document processing."
        )
        self.doc = Document()
        self.image_cache = {}  # Reset image cache for this run
        logger.debug("Minifying HTML content...")
        html_content = minify(
            html_content,
            allow_noncompliant_unquoted_attribute_values=False,
            allow_optimal_entities=False,
            allow_removing_spaces_between_attributes=True,
            keep_closing_tags=True,
            keep_comments=False,
            keep_html_and_head_opening_tags=True,
            keep_input_type_text_attr=False,
            keep_ssi_comments=False,
            minify_css=True,
            minify_doctype=False,
            minify_js=True,
            preserve_brace_template_syntax=False,
            preserve_chevron_percent_template_syntax=False,
            remove_bangs=False,
            remove_processing_instructions=False,
        )
        try:
            soup = BeautifulSoup(html_content, "lxml")
        except FeatureNotFound:
            soup = BeautifulSoup(html_content, "html.parser")
        except Exception as e:
            logger.error("BeautifulSoup parsing error: %s", e)
            return
        logger.debug("HTML parsing complete. Decomposing script tags.")

        # Strip all <script> tags
        for script_tag in soup.find_all("script"):
            logger.debug("Removing script tag: %s", str(script_tag)[:100])
            script_tag.decompose()

        body = soup.find("body")
        process_root = body if body else soup

        for element in process_root.children:  # type: ignore
            if isinstance(element, (Tag, NavigableString)):
                logger.debug("Processing root child: %s", str(element)[:100])
                self._process_block_element(element, self.doc, [])

        try:
            logger.debug(
                "Attempting to save DOCX file to: %s", self.output_path
            )
            self.doc.save(self.output_path)
            logger.info("DOCX file successfully saved to %s", self.output_path)
            QDesktopServices.openUrl(QUrl.fromLocalFile(self.output_path))
        except Exception as e:
            logger.error(
                "Error saving DOCX file to %s: %s", self.output_path, e
            )

    def _run_js_async(self, js_code: str, callback) -> Any:
        page = self.view.page()
        assert page is not None
        page.runJavaScript(js_code, callback)

    def _add_element_as_image(self, element_id: str, parent_docx_object):
        # Use self.elements_map to get geometry
        if not self.full_page_qimage:
            logger.error(
                "Missing full page screenshot for element %s", element_id
            )
            return
        logger.debug("Adding element %s as image.", element_id)

        element_data = self.elements_map.get(element_id)
        if not element_data:
            logger.error("No map data found for element %s", element_id)
            return

        geometry_data = element_data.get("geometry")
        if not geometry_data:
            logger.error(
                "Missing geometry for image/svg element %s", element_id
            )
            return

        if element_id in self.image_cache:
            img_bytes = self.image_cache[element_id]
        else:
            logger.debug("Element %s not in image cache, cropping.", element_id)
            # Use geometry_data directly for coordinates and dimensions
            x = int(geometry_data["docX"])
            y = int(geometry_data["docY"])
            width = int(geometry_data["docWidth"])
            height = int(geometry_data["docHeight"])

            if width <= 0 or height <= 0:
                logger.error(
                    "Invalid dimensions for element %s: w=%s, h=%s",
                    element_id,
                    width,
                    height,
                )
                return

            cropped_qimage = self.full_page_qimage.copy(x, y, width, height)
            if cropped_qimage.isNull():
                logger.error("Failed to crop image for element %s", element_id)
                return
            logger.debug("Successfully cropped image for %s.", element_id)

            byte_array = QByteArray()
            buffer = QBuffer(byte_array)
            buffer.open(QIODevice.OpenModeFlag.WriteOnly)
            cropped_qimage.save(buffer, "PNG")
            img_bytes = bytes(byte_array.data())
            self.image_cache[element_id] = img_bytes

        if img_bytes:
            img_stream = io.BytesIO(img_bytes)
            # Use geometry_data directly for width/height
            w_px = geometry_data["docWidth"]
            h_px = geometry_data["docHeight"]

            para_for_img = None
            if hasattr(parent_docx_object, "add_paragraph") and not isinstance(
                parent_docx_object, DocumentObject
            ):
                # If parent is a cell or similar that can add paragraphs
                para_for_img = parent_docx_object.add_paragraph()
            elif hasattr(
                parent_docx_object, "add_run"
            ):  # Parent is already a paragraph
                para_for_img = parent_docx_object
            else:  # Document object or fallback
                para_for_img = self.doc.add_paragraph()

            try:
                dpi = 96.0
                w_inch = Inches(w_px / dpi) if w_px > 0 else None
                h_inch = Inches(h_px / dpi) if h_px > 0 else None

                run_for_picture = para_for_img.add_run()  # type: ignore
                if w_inch and h_inch:
                    run_for_picture.add_picture(
                        img_stream, width=w_inch, height=h_inch
                    )
                elif w_inch:
                    run_for_picture.add_picture(img_stream, width=w_inch)
                elif h_inch:
                    run_for_picture.add_picture(img_stream, height=h_inch)
                else:  # Add with original pixel size
                    run_for_picture.add_picture(img_stream)
            except Exception as e:
                logger.error(
                    "Error adding cropped image %s to DOCX: %s",
                    element_id,
                    e,
                )
            logger.debug("Image %s added to DOCX.", element_id)

    # --- Color and Style Parsers (largely unchanged but may need review) ---
    def _hex_to_rgb(self, hex_color: str) -> RGBColor | None:
        hex_color = hex_color.lstrip("#")
        if len(hex_color) == 3:
            hex_color = "".join([c * 2 for c in hex_color])
        if len(hex_color) == 6:
            try:
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                return RGBColor(r, g, b)
            except ValueError:
                return None
        return None

    def _get_color(self, css_color_string: str) -> RGBColor | None:
        css_color_string = css_color_string.lower().strip()
        named_colors = {
            "black": RGBColor(0, 0, 0),
            "white": RGBColor(255, 255, 255),
            "red": RGBColor(255, 0, 0),
            "green": RGBColor(0, 128, 0),
            "blue": RGBColor(0, 0, 255),
            "yellow": RGBColor(255, 255, 0),
            "transparent": None,
        }
        if css_color_string in named_colors:
            return named_colors[css_color_string]
        if css_color_string.startswith("#"):
            return self._hex_to_rgb(css_color_string)
        elif css_color_string.startswith("rgb("):
            try:
                match = re.match(
                    r"rgb\\((\\d+),\\s*(\\d+),\\s*(\\d+)\\)", css_color_string
                )
                if match:
                    r, g, b = map(int, match.groups())
                    return RGBColor(r, g, b)
            except Exception:
                return None
        return None

    def _parse_styles(self, element: Tag) -> tuple[dict[str, Any], list[str]]:
        styles: dict[str, Any] = {}
        style_attr = element.get("style")
        if style_attr and isinstance(style_attr, str):
            for style_declaration in style_attr.split(";"):
                style_declaration = style_declaration.strip()
                if ":" in style_declaration:
                    prop, val = style_declaration.split(":", 1)
                    prop = prop.strip().lower()
                    val = val.strip()
                    # logger.debug(f"Parsed style from element
                    #           {element.name}: {prop}={val}")
                    if prop == "--bs-border-opacity":
                        try:
                            styles["border_opacity"] = float(val)
                        except ValueError:
                            logger.warning("Invalid opacity: %s", val)
                    else:
                        styles[prop] = val

        final_class_list: list[str] = []
        class_attr = element.get("class")
        if isinstance(class_attr, str):
            final_class_list = class_attr.split()
        elif isinstance(class_attr, list):
            final_class_list = [
                str(c) for c in class_attr if isinstance(c, str)
            ]
        # if final_class_list:
        #     logger.debug(f"Parsed classes from element
        #           {element.name}: {final_class_list}")
        return styles, final_class_list

    # --- Content Processors (adapted for new image handling) ---
    def _apply_formatting_to_run(self, run, element_tag: Tag):
        tag_name = element_tag.name.lower()
        styles, _ = self._parse_styles(element_tag)

        if tag_name in ["b", "strong"] or styles.get("font-weight") == "bold":
            run.bold = True
        if tag_name in ["i", "em"] or styles.get("font-style") == "italic":
            run.italic = True

        text_decoration = styles.get("text-decoration", "")
        if tag_name == "u" or "underline" in text_decoration:
            run.underline = True
        if (
            tag_name in ["s", "strike", "del"]
            or "line-through" in text_decoration
        ):
            run.strike = True

        color_str = styles.get("color")
        if color_str:
            color = self._get_color(color_str)
            if color:
                logger.debug(
                    "Applying font color %s to run for element %s",
                    color,
                    element_tag.name,
                )
                run.font.color.rgb = color

        font_family_str = styles.get("font-family")
        if font_family_str:
            primary_font = (
                font_family_str.split(",")[0]
                .strip()
                .replace("'", "")
                .replace('"', "")
            )
            if primary_font:
                run.font.name = primary_font

        font_size_str = styles.get("font-size")
        if font_size_str:
            try:
                val_str = re.sub(r"[^\\d\\.]", "", font_size_str)
                if val_str:
                    val = float(val_str)
                    if "pt" in font_size_str:
                        run.font.size = Pt(val)
                    elif "px" in font_size_str:
                        run.font.size = Pt(val * 0.75)
            except ValueError:
                pass

        if tag_name == "a":
            if not run.font.color.rgb:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                logger.debug("Applied default link color to run for <a> tag")
            if not run.underline:
                run.underline = WD_UNDERLINE.SINGLE
                logger.debug(
                    "Applied default link underline to run for <a> tag"
                )

    def _process_inline_content(
        self,
        html_node: PageElement,
        docx_parent_paragraph,
        active_format_tags: list[Tag],
        first: bool = False,
        is_dt_content: bool = False,
        crt_par_classes: list[str] | None = None,
    ):
        if isinstance(html_node, NavigableString):
            text = str(html_node)
            if len(text) > 0:
                prefix = " " if text[0].isspace() and not first else ""
                suffix = " " if text[-1].isspace() else ""
                text = prefix + text.strip() + suffix
                logger.debug(
                    "Processing NavigableString: '%s' (first=%s, is_dt=%s)",
                    text[:50].replace("\n", " "),
                    first,
                    is_dt_content,
                )

                run = docx_parent_paragraph.add_run(text)
                if is_dt_content:
                    run.bold = True
                for fmt_tag in active_format_tags:
                    self._apply_formatting_to_run(run, fmt_tag)

        elif isinstance(html_node, Tag):
            # Check if element should be skipped based on d-none
            # Safely get element_id for map lookup
            raw_element_id = html_node.get("data-docgen-id")
            element_id_str: str | None = None
            if isinstance(raw_element_id, list):
                if raw_element_id:  # Check if list is not empty
                    element_id_str = str(raw_element_id[0])
            elif isinstance(raw_element_id, str):
                element_id_str = raw_element_id

            if element_id_str:
                element_data = self.elements_map.get(element_id_str)
                if element_data and element_data.get("is_hidden"):
                    logger.debug(
                        "Skip hidden inline %s (%s)",
                        element_id_str,
                        html_node.name,
                    )
                    return  # Skip this hidden element

            tag_name = html_node.name.lower()
            logger.debug("Processing inline tag: <%s>", tag_name)

            # Ignore <button> tags and their content
            if tag_name == "button":
                logger.debug(
                    "Ignoring <button> tag and its content in block "
                    "context: %s",
                    str(html_node)[:50],
                )
                return

            if tag_name == "br":
                docx_parent_paragraph.add_run().add_break()
                return
            logger.debug("Added <br> to paragraph.")

            if tag_name == "sup":
                sup_text = (
                    html_node.get_text().strip()
                )  # Get all text within <sup>
                if sup_text:  # Only add run if there's text
                    logger.debug("Applying <sup> with text: '%s'", sup_text)
                    run = docx_parent_paragraph.add_run(sup_text)
                    run.font.superscript = True
                    # Apply any inherited formatting from parent tags
                    for fmt_tag in active_format_tags:
                        self._apply_formatting_to_run(run, fmt_tag)
                    # Apply any direct styling on the <sup> tag itself
                    self._apply_formatting_to_run(run, html_node)
                return  # Consumed <sup> tag and its content
            logger.debug("Finished processing <sup> tag.")

            if tag_name in ["img", "svg"]:
                # Safely get element_id for map lookup
                raw_img_id = html_node.get("data-docgen-id")
                img_id_str: str | None = None
                if isinstance(raw_img_id, list):
                    if raw_img_id:  # Check if list is not empty
                        img_id_str = str(raw_img_id[0])
                elif isinstance(raw_img_id, str):
                    img_id_str = raw_img_id

                if img_id_str:
                    element_map_data = self.elements_map.get(img_id_str, {})
                    if element_map_data.get("geometry"):
                        self._add_element_as_image(
                            img_id_str, docx_parent_paragraph
                        )
                    else:
                        logger.debug(
                            f"Skip inline img/svg {img_id_str}, no geometry."
                        )
                else:
                    logger.error(
                        "Skip inline <%s> no data-docgen-id: %.30s",
                        tag_name,
                        str(html_node),
                    )
                return  # Consumed img/svg (or skipped)
            logger.debug("Finished processing <img> or <svg> tag.")

            current_active_tags = list(active_format_tags)
            is_inline_formatting_provider = tag_name in [
                "b",
                "strong",
                "i",
                "em",
                "u",
                "s",
                "strike",
                "del",
                "span",
                "font",
                "a",
            ]
            if is_inline_formatting_provider:
                current_active_tags.append(html_node)

            if tag_name == "a":
                href = html_node.get("href")
                link_text = "".join(
                    map(str, html_node.find_all(string=True, recursive=True))
                ).strip()  # Ensure strings
                if href and link_text:
                    try:
                        h_run = docx_parent_paragraph.add_hyperlink(
                            href, link_text, is_external=True
                        )
                        for fmt_tag in current_active_tags:
                            self._apply_formatting_to_run(h_run, fmt_tag)
                    except Exception as e:
                        logger.error(
                            "Failed to add hyperlink for %s: %s. "
                            "Adding as text.",
                            href,
                            e,
                        )
                        for i, child in enumerate(html_node.children):
                            self._process_inline_content(
                                child,
                                docx_parent_paragraph,
                                current_active_tags,
                                first=(i == 0),
                                is_dt_content=False,
                                crt_par_classes=crt_par_classes,
                            )
                    return
                # Fallback for 'a' if no href/text or error
                for i, child in enumerate(html_node.children):
                    self._process_inline_content(
                        child,
                        docx_parent_paragraph,
                        current_active_tags,
                        first=(i == 0),
                        is_dt_content=False,
                        crt_par_classes=crt_par_classes,
                    )
            else:
                for i, child in enumerate(html_node.children):
                    self._process_inline_content(
                        child,
                        docx_parent_paragraph,
                        current_active_tags,
                        first=(i == 0),
                        is_dt_content=False,
                        crt_par_classes=crt_par_classes,
                    )

    # --- Table Handlers (largely unchanged but verify context for
    # image processing if any) ---
    def _set_cell_shading(self, cell, color_str: str):
        color = self._get_color(color_str)
        if color and color_str.lower() != "transparent":
            try:
                logger.debug("Applying cell shading %s", color_str)
                shd = OxmlElement("w:shd")
                # RGBColor.__str__ already returns correct hex format for Word
                # XML
                shd.set(qn("w:fill"), str(color))
                shd.set(qn("w:val"), "clear")
                tcPr = cell._tc.get_or_add_tcPr()
                for existing_shd in tcPr.xpath("./w:shd"):
                    tcPr.remove(existing_shd)
                tcPr.append(shd)
            except Exception as e:
                logger.error(
                    "Error applying cell shading %s: %s",
                    color_str,
                    e,
                )

    def _set_cell_border_color(
        self,
        border_side_element: OxmlElementType,
        color_rgb: tuple[int, int, int],  # Expect (R,G,B) tuple
        size_pt: int = 4,  # This is w:sz unit (eighths of a point)
        alpha: float = 1.0,  # Opacity (0.0 to 1.0)
    ):
        r, g, b = color_rgb
        actual_r, actual_g, actual_b = r, g, b

        if alpha < 1.0 and alpha >= 0.0:  # Apply opacity by blending with white
            actual_r = int(r * alpha + 255 * (1 - alpha))
            actual_g = int(g * alpha + 255 * (1 - alpha))
            actual_b = int(b * alpha + 255 * (1 - alpha))

        # Ensure values are within valid range
        final_r = max(0, min(actual_r, 255))
        final_g = max(0, min(actual_g, 255))
        final_b = max(0, min(actual_b, 255))

        # Convert to hex format expected by Word XML (without # prefix)
        hex_color = f"{final_r:02X}{final_g:02X}{final_b:02X}"

        border_side_element.set(qn("w:val"), "single")
        border_side_element.set(qn("w:sz"), str(size_pt))
        border_side_element.set(qn("w:color"), hex_color)

    def _apply_cell_borders(
        self,
        cell,
        cell_styles: dict,
        cell_classes: list[str],
        table_classes: list[str],
    ):
        tcPr = cell._tc.get_or_add_tcPr()
        # type: ignore # mypy issue with lxml-based find
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
            logger.debug("Created new w:tcBorders for cell.")
        else:
            logger.debug(
                "Found existing w:tcBorders for cell. "
                "Clearing existing border elements..."
            )

        for border_tag_to_clear in [
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ]:
            existing = tcBorders.find(  # type: ignore
                qn(f"w:{border_tag_to_clear}")
            )
            if existing is not None:
                tcBorders.remove(existing)

        base_r, base_g, base_b = DEFAULT_BORDER_COLOR_RGB
        logger.debug(
            "Applying borders. Table classes: %s, Cell classes: %s, "
            "Cell styles: %s",
            table_classes,
            cell_classes,
            cell_styles,
        )

        if (
            "border-success" in cell_classes
            or "border-success" in table_classes
        ):
            base_r, base_g, base_b = BORDER_COLOR_SUCCESS_RGB

        border_opacity = cell_styles.get("border_opacity", 1.0)
        if not isinstance(border_opacity, (float, int)) or not (
            0.0 <= border_opacity <= 1.0
        ):
            border_opacity = 1.0

        sides_to_border = ["top", "bottom", "left", "right"]
        apply_default_bordered = "table-bordered" in table_classes

        for side in sides_to_border:
            border_definition_from_style = cell_styles.get(f"border-{side}")
            apply_this_side = False
            parsed_width_val = DEFAULT_BORDER_WIDTH_PT

            # If table-bordered, always apply border for the side,
            # ignoring cell-specific border-style unless it explicitly says
            # none.
            if apply_default_bordered:
                apply_this_side = True
                # If cell explicitly sets this border to none, respect that
                # even if table is bordered
                if border_definition_from_style and (
                    "none" in border_definition_from_style
                    or "0px" in border_definition_from_style
                    or "0pt" in border_definition_from_style
                ):
                    apply_this_side = False
            elif border_definition_from_style:
                # Not table-bordered, but cell has a specific border style
                if (
                    "none" in border_definition_from_style
                    or "0px" in border_definition_from_style
                    or "0pt" in border_definition_from_style
                ):
                    apply_this_side = False
                else:
                    apply_this_side = True
                    # TODO: More detailed parsing of border-{side} for width,
                    # color etc. For now, if style exists and is not none, it
                    # uses default width/calculated color.

            if apply_this_side:
                if border_opacity < 0.05:  # Effectively transparent, skip
                    continue

                border_el = OxmlElement(f"w:{side}")
                border_sz = max(1, int(parsed_width_val * 8))
                logger.debug(
                    "Setting border %s: color_rgb=(%s,%s,%s), "
                    "size_pt=%s, alpha=%s",
                    side,
                    base_r,
                    base_g,
                    base_b,
                    border_sz,
                    border_opacity,
                )
                self._set_cell_border_color(
                    border_el,
                    (base_r, base_g, base_b),
                    size_pt=border_sz,
                    alpha=border_opacity,
                )
                tcBorders.append(border_el)
                logger.debug("Appended %s border element to tcBorders.", side)

    @define
    class TableData:
        """Keeps track of the table data as we process the HTML.

        Attributes:
            html_grid: The grid for the docx table.
            html_rows: The rows from the HTML table.
            max_cols: The maximum number of columns across all rows.
            table_styles: The styles for the table.
            table_classes: The classes for the table.
            grid_r_idx: The current row index in the doc table.
        """

        cv: "HtmlToDocxConverter"
        html_grid: List[List[Tag | str | None]] = field(factory=list)
        html_rows: List[Tag] = field(factory=list)
        max_cols: int = field(default=0)
        table_styles: dict[str, Any] = field(factory=dict)
        table_classes: list[str] = field(factory=list)
        grid_r_idx: int = -1
        doc_table: Any = None
        num_logical_rows: int = 0

        def _collect_table_rows(self, table_element: Tag):
            """Collect all rows from a table element."""

            # Collect rows from the table header.
            thead = table_element.find("thead", recursive=False)
            if thead and isinstance(thead, Tag):
                for tag in thead.find_all("tr", recursive=False):
                    if isinstance(tag, Tag):
                        self.html_rows.append(tag)

            # Collect rows from the table body.
            tbody_elements = table_element.find_all("tbody", recursive=False)
            if tbody_elements:
                for tbody in tbody_elements:
                    if isinstance(tbody, Tag):
                        for tag in tbody.find_all("tr", recursive=False):
                            if isinstance(tag, Tag):
                                self.html_rows.append(tag)

            # Collect rows from within the table directly.
            for tag in table_element.find_all("tr", recursive=False):
                if isinstance(tag, Tag):
                    self.html_rows.append(tag)

            # Collect rows from the table footer.
            tfoot = table_element.find("tfoot", recursive=False)
            if tfoot and isinstance(tfoot, Tag):
                for tag in tfoot.find_all("tr", recursive=False):
                    if isinstance(tag, Tag):
                        self.html_rows.append(tag)

            logger.debug(
                "Collected %d HTML rows for the table.", len(self.html_rows)
            )

        def _handle_table_row(self, r_idx, tr_element_maybe_str) -> None:
            if not isinstance(tr_element_maybe_str, Tag):
                logger.debug(
                    "Skipping non-<tr> element at index %d: %s",
                    r_idx,
                    str(tr_element_maybe_str)[:50],
                )
                return

            tr_element: Tag = tr_element_maybe_str

            # Check if the <tr> itself is hidden

            # This is the ID that we created ourselves in javascript.
            raw_tr_id = tr_element.get("data-docgen-id")
            tr_id_str: str | None = None
            if isinstance(raw_tr_id, list):
                if raw_tr_id:
                    tr_id_str = str(raw_tr_id[0])
            elif isinstance(raw_tr_id, str):
                tr_id_str = raw_tr_id
            else:
                logger.warning(
                    "Unexpected type for tr_id '%s': %s",
                    raw_tr_id,
                    type(raw_tr_id),
                )

            # Skip this entire row if the row is hidden.
            if tr_id_str:
                tr_element_data = self.cv.elements_map.get(tr_id_str)
                if tr_element_data and tr_element_data.get("is_hidden"):
                    logger.debug(
                        "Skipping hidden <tr> element with ID %s", tr_id_str
                    )
                    return

            # Increment grid_r_idx and append row only for visible <tr>
            self.grid_r_idx += 1
            self.html_grid.append([])
            current_col_idx = 0
            if not isinstance(tr_element, Tag):
                # Should have been caught by first check
                logger.error(
                    "Unexpected type for tr_element at index %d: %s",
                    r_idx,
                    type(tr_element),
                )
                return

            # Process each cell in the row.
            for td_th_element in tr_element.find_all(
                ["td", "th"], recursive=False
            ):
                if not isinstance(td_th_element, Tag):
                    logger.warning(
                        "Skipping non-<td>/<th> element at index %d: %s",
                        r_idx,
                        str(td_th_element)[:50],
                    )
                    return

                # Move to the next available column in the grid
                while (
                    len(self.html_grid[self.grid_r_idx]) > current_col_idx
                    and self.html_grid[self.grid_r_idx][current_col_idx]
                    is not None
                ):
                    current_col_idx += 1
                logger.debug(
                    "  Cell processing: grid_r_idx=%d, current_col_idx=%d",
                    self.grid_r_idx,
                    current_col_idx,
                )

                # Get colspan and rowspan attributes
                colspan = self.cv._get_attribute_as_int(
                    td_th_element, "colspan", 1
                )
                rowspan = self.cv._get_attribute_as_int(
                    td_th_element, "rowspan", 1
                )
                logger.debug(
                    "    HTML Cell (%s): colspan=%d, rowspan=%d",
                    str(td_th_element.name),
                    colspan,
                    rowspan,
                )

                # Process rowspan
                for i in range(rowspan):
                    target_r_in_grid = self.grid_r_idx + i

                    # Ensure rows up to the target row exist
                    self.html_grid.extend(
                        [
                            []
                            for _ in range(
                                len(self.html_grid), target_r_in_grid + 1
                            )
                        ]
                    )

                    # Ensure enough columns in the target row
                    while (
                        len(self.html_grid[target_r_in_grid])
                        < current_col_idx + colspan
                    ):
                        self.html_grid[target_r_in_grid].append(None)

                    # Assign cell or mark as SPAN
                    for j in range(colspan):
                        cell_value = (
                            td_th_element if i == 0 and j == 0 else "SPAN"
                        )
                        self.html_grid[target_r_in_grid][
                            current_col_idx + j
                        ] = cell_value

                # Update column index for the next cell
                current_col_idx += colspan

                # Track maximum number of columns across all rows.
                self.max_cols = max(self.max_cols, current_col_idx)

        def _merged_cell(
            self,
            doc_cell,
            r_idx_grid_local: int,
            c_idx: int,
            colspan: int,
            rowspan: int,
        ) -> None:
            end_r_idx = min(
                r_idx_grid_local + rowspan - 1, self.num_logical_rows - 1
            )
            end_c_idx = min(c_idx + colspan - 1, self.max_cols - 1)
            if end_r_idx > r_idx_grid_local or end_c_idx > c_idx:
                try:
                    doc_cell.merge(self.doc_table.cell(end_r_idx, end_c_idx))
                    logger.debug(
                        "Merged cell at (grid_row=%d, col=%d) to "
                        "(grid_row=%d, col=%d)",
                        r_idx_grid_local,
                        c_idx,
                        end_r_idx,
                        end_c_idx,
                    )
                except Exception as e:
                    logger.warning(f"Cell merge failed: {e}")

        def _copy_borders(
            self,
            primary_tcBorders: Any,
            continued_tcPr: Any,
            r_idx_grid_local: int,
            c_idx: int,
            rowspan: int,
        ):
            """Copy borders from primary cell to continued cells"""
            for row_offset in range(1, rowspan):
                true_row_offset = r_idx_grid_local + row_offset
                if true_row_offset < self.num_logical_rows:
                    continued_cell = self.doc_table.cell(true_row_offset, c_idx)
                    continued_tcPr = continued_cell._tc.get_or_add_tcPr()

                    logger.debug(
                        "    Setting w:vMerge from restart to "
                        "continue for cell (grid_row=%d, col=%d)",
                        r_idx_grid_local + row_offset,
                        c_idx,
                    )

                    # Set vMerge to "continue" for continued cells
                    vmerge_elem = continued_tcPr.find(qn("w:vMerge"))
                    if vmerge_elem is None:
                        vmerge_elem = OxmlElement("w:vMerge")
                        continued_tcPr.append(vmerge_elem)
                    vmerge_elem.set(qn("w:val"), "continue")
                    logger.debug(
                        "     Changed w:vMerge from restart to "
                        "continue for cell (grid_row=%d, col=%d)",
                        true_row_offset,
                        c_idx,
                    )

                    # Ensure all four sides (top, bottom, left, right) are
                    # cloned. Remove any existing <w:tcBorders> child
                    for old_borders in continued_tcPr.xpath("./w:tcBorders"):
                        continued_tcPr.remove(old_borders)

                    # Create a new <w:tcBorders> and explicitly clone top,
                    # bottom, left, right
                    new_tc_borders_el = OxmlElement("w:tcBorders")
                    # The four possible side tags, in Word’s expected order
                    for side_tag in ("top", "bottom", "left", "right"):
                        # look for that side under primary_tcBorders
                        primary_side_el = primary_tcBorders.find(
                            qn(f"w:{side_tag}")
                        )
                        if primary_side_el is not None:
                            cloned_side_el = deepcopy(primary_side_el)
                            new_tc_borders_el.append(cloned_side_el)
                    continued_tcPr.append(new_tc_borders_el)

        def _handle_cell(self, r_idx_grid_local: int, c_idx) -> None:
            html_cell_content = self.html_grid[r_idx_grid_local][c_idx]
            if html_cell_content is None or html_cell_content == "SPAN":
                return

            if not isinstance(html_cell_content, Tag):
                return

            html_cell_element: Tag = html_cell_content
            logger.debug(
                "Processing doc_cell at (grid_row=%d, col=%d) for "
                "HTML cell: %s",
                r_idx_grid_local,
                c_idx,
                str(html_cell_element)[:50],
            )
            doc_cell = self.doc_table.cell(r_idx_grid_local, c_idx)

            colspan = self.cv._get_attribute_as_int(
                html_cell_element, "colspan", 1
            )
            rowspan = self.cv._get_attribute_as_int(
                html_cell_element, "rowspan", 1
            )

            if rowspan > 1 or colspan > 1:
                self._merged_cell(
                    doc_cell,
                    r_idx_grid_local,
                    c_idx,
                    colspan,
                    rowspan,
                )

            # Only remove the default <w:p> if we actually have content to add
            # i.e. if the HTML <td>/<th> has at least one child (text or tag)
            if html_cell_element is not None and list(
                html_cell_element.children
            ):
                if doc_cell.paragraphs and doc_cell.paragraphs[0].text == "":
                    p_element = doc_cell.paragraphs[0]._element
                    p_element.getparent().remove(p_element)

            cell_styles, cell_classes = self.cv._parse_styles(html_cell_element)

            # Table striping and background color
            is_striped_table = "table-striped" in self.table_classes
            if is_striped_table and (r_idx_grid_local % 2 != 0):
                if not cell_styles.get("background-color"):
                    shade_color = RGBColor(*TABLE_STRIPED_BG_COLOR_RGB)
                    self.cv._set_cell_shading(doc_cell, str(shade_color))
            bg_color_str = cell_styles.get("background-color")
            if bg_color_str:
                self.cv._set_cell_shading(doc_cell, bg_color_str)

            # Apply borders to the primary cell
            self.cv._apply_cell_borders(
                doc_cell, cell_styles, cell_classes, self.table_classes
            )

            # Populate the cell with content from html_cell_element
            if html_cell_element and hasattr(html_cell_element, "children"):
                active_tags_for_cell_content = [html_cell_element]
                for child_node in html_cell_element.children:
                    self.cv._process_block_element(
                        child_node, doc_cell, active_tags_for_cell_content
                    )

            # Direct OXML tcBorders copy for continued cells in rowspan
            if rowspan > 1:
                primary_tcPr = doc_cell._tc.get_or_add_tcPr()
                primary_tcBorders = primary_tcPr.find(
                    qn("w:tcBorders")
                )  # type: ignore[arg-type]
                logger.debug(
                    "  Rowspan > 1 for cell at (grid_row=%d, col=%d). "
                    "Primary tcBorders exists: %s",
                    r_idx_grid_local,
                    c_idx,
                    primary_tcBorders is not None,
                )

                if primary_tcBorders is not None:
                    self._copy_borders(
                        primary_tcBorders,
                        doc_cell._tc.get_or_add_tcPr(),
                        r_idx_grid_local,
                        c_idx,
                        rowspan,
                    )

    def _handle_table(self, table_element: Tag, parent_docx_object):
        """Process a table element and convert it to a Docx table."""
        logger.debug(
            "--- Starting _handle_table for: %s ---", str(table_element)[:100]
        )
        table_data = self.TableData(cv=self)

        # Get table classes
        table_data.table_styles, table_data.table_classes = self._parse_styles(
            table_element
        )
        logger.debug(
            "Table styles: %s, Table classes: %s",
            table_data.table_styles,
            table_data.table_classes,
        )

        # Collect html rows.
        table_data._collect_table_rows(table_element)

        # Process each row.
        for r_idx, tr_element_maybe_str in enumerate(table_data.html_rows):
            table_data._handle_table_row(r_idx, tr_element_maybe_str)

        # Ensure the grid has enough columns.
        table_data.num_logical_rows = len(table_data.html_grid)
        for r_list in table_data.html_grid:
            while len(r_list) < table_data.max_cols:
                r_list.append(None)

        # If the table has 0 rows or 0 columns, skip it.
        if table_data.num_logical_rows == 0 or table_data.max_cols == 0:
            logger.warning(
                "Table has 0 rows or 0 columns after processing grid. "
                "Skipping table."
            )
            return

        table_data.doc_table = parent_docx_object.add_table(
            rows=table_data.num_logical_rows,
            cols=table_data.max_cols,
        )
        # Apply table-level styles like table-layout: fixed if needed (not
        # requested yet) doc_table.autofit = False doc_table.layout_type =
        # WD_TABLE_LAYOUT.FIXED

        for r_idx_grid_local in range(
            table_data.num_logical_rows
        ):  # Iterate using num_logical_rows (len of html_grid)
            for c_idx in range(table_data.max_cols):
                table_data._handle_cell(r_idx_grid_local, c_idx)

        logger.debug(
            "--- Finished _handle_table for: %s ---", str(table_element)[:100]
        )

    def _process_block_element(
        self,
        element: PageElement,
        parent_docx_object,
        active_format_tags: list[Tag],
        indent_level_inches: float | None = None,
        crt_par_classes: list[str] | None = None,
    ):
        current_paragraph: Any = None

        if isinstance(element, NavigableString):
            # Delegate NavigableString to _process_inline_content
            # Ensure a paragraph context exists in parent_docx_object
            target_p_for_nav_str = None
            is_first_run = True
            logger.debug(
                "Processing NavigableString (block context): '%s'",
                str(element)[:50].replace("\n", " "),
            )
            if hasattr(parent_docx_object, "_tc"):  # Parent is a cell
                if parent_docx_object.paragraphs:
                    target_p_for_nav_str = parent_docx_object.paragraphs[-1]
                    is_first_run = not bool(target_p_for_nav_str.runs)
                else:
                    target_p_for_nav_str = parent_docx_object.add_paragraph()
            elif hasattr(
                parent_docx_object, "add_run"
            ):  # Parent is already a paragraph
                target_p_for_nav_str = parent_docx_object
                is_first_run = not bool(target_p_for_nav_str.runs)
            else:  # Fallback (e.g., document root)
                target_p_for_nav_str = self._create_paragraph_for_block(
                    parent_docx_object, indent_level_inches
                )

            self._process_inline_content(
                element,  # The NavigableString itself
                target_p_for_nav_str,
                active_format_tags,
                first=is_first_run,  # Pass 'first' status
                # is_dt_content and crt_par_classes ar e
                # context-dependent
            )
            return

        if not isinstance(element, Tag):
            return  # Ignore other PageElement types like Comment

        # Check if element should be skipped based on d-none
        # Safely get element_id for map lookup
        raw_element_id_block = element.get("data-docgen-id")
        element_id_block_str: str | None = None
        if isinstance(raw_element_id_block, list):
            if raw_element_id_block:  # Check if list is not empty
                element_id_block_str = str(raw_element_id_block[0])
        elif isinstance(raw_element_id_block, str):
            element_id_block_str = raw_element_id_block

        if element_id_block_str:
            element_data = self.elements_map.get(element_id_block_str)
            if element_data and element_data.get("is_hidden"):
                logger.debug(
                    f"Skip hidden block {element_id_block_str} ({element.name})"
                )
                return  # Skip this hidden element

        tag_name = element.name.lower()
        logger.debug(
            "Processing block tag: <%s> with indent %s",
            tag_name,
            indent_level_inches,
        )

        # Ignore <button> tags and their content
        if tag_name == "button":
            logger.debug(
                "Ignoring <button> tag and its content in block " "context: %s",
                str(element)[:50],
            )
            return

        if tag_name in ["img", "svg"]:
            # Safely get element_id for map lookup
            raw_img_id_block = element.get("data-docgen-id")
            img_id_block_str: str | None = None
            if isinstance(raw_img_id_block, list):
                if raw_img_id_block:  # Check if list is not empty
                    img_id_block_str = str(raw_img_id_block[0])
            elif isinstance(raw_img_id_block, str):
                img_id_block_str = raw_img_id_block

            if img_id_block_str:
                element_map_data = self.elements_map.get(img_id_block_str, {})
                if element_map_data.get("geometry"):
                    target_for_image = None
                    if hasattr(
                        parent_docx_object, "add_paragraph"
                    ) and not isinstance(parent_docx_object, DocumentObject):
                        logger.debug(
                            "Image parent is a cell/similar, adding new "
                            "paragraph for image."
                        )
                        target_for_image = parent_docx_object
                    elif hasattr(parent_docx_object, "add_run"):
                        logger.debug("Image parent is a paragraph.")
                        target_for_image = parent_docx_object
                    else:
                        target_for_image = self.doc
                    self._add_element_as_image(
                        img_id_block_str, target_for_image
                    )
                else:
                    logger.debug(
                        f"Skip block img/svg {img_id_block_str}, no geometry."
                    )
            else:
                logger.error(
                    "Skip block <%s> no data-docgen-id: %.30s",
                    tag_name,
                    str(element),
                )
            return  # Consumed img/svg or skipped

        # Handle p, div, h* separately from other block/inline logic
        if tag_name in ["p", "div", "h1", "h2", "h3", "h4", "h5", "h6"]:
            current_paragraph = self._create_paragraph_for_block(
                parent_docx_object, indent_level_inches
            )

            if tag_name.startswith("h"):
                try:
                    current_paragraph.style = f"Heading {int(tag_name[1])}"
                except (ValueError, IndexError):
                    logger.warning(
                        "Could not apply heading style for %s", tag_name
                    )
                logger.debug(
                    "Applied style '%s' to paragraph for <%s>",
                    current_paragraph.style.name,  # type: ignore
                    tag_name,
                )

            # Get styles and classes for the block element itself
            element_styles, element_classes = self._parse_styles(element)
            new_active_tags = list(active_format_tags)
            # DIVs can provide formatting context via styles/classes
            if tag_name == "div":
                new_active_tags.append(element)
                # Still useful for style inheritance if not using classes
                # directly here

            for i, child in enumerate(element.children):
                self._process_inline_content(
                    child,
                    current_paragraph,
                    new_active_tags,
                    first=(i == 0),
                    is_dt_content=False,
                    crt_par_classes=element_classes,
                )

        elif tag_name in ["ul", "ol"]:
            for item in element.find_all("li", recursive=False):
                if not isinstance(item, Tag):
                    continue
                style = "ListBullet" if tag_name == "ul" else "ListNumber"
                logger.debug("Adding <li> item with style: %s", style)
                # Content of <li> processed into this new paragraph
                li_paragraph = self.doc.add_paragraph(style=style)
                if indent_level_inches:
                    li_paragraph.paragraph_format.left_indent = Inches(
                        indent_level_inches
                    )
                for i, child_content in enumerate(item.children):
                    self._process_inline_content(
                        child_content,
                        li_paragraph,
                        active_format_tags,
                        first=(i == 0),
                        is_dt_content=False,
                        crt_par_classes=crt_par_classes,
                    )

        elif tag_name == "dl":
            self._handle_dl_element(
                element,
                parent_docx_object,
                active_format_tags,
                indent_level_inches,
                crt_par_classes=crt_par_classes,
            )
            logger.debug("Finished processing <dl> element.")

        elif tag_name == "table":
            # Pass table classes from _parse_styles to _handle_table
            self._handle_table(element, parent_docx_object)

        elif tag_name == "hr":
            p = self.doc.add_paragraph()
            if indent_level_inches:
                p.paragraph_format.left_indent = Inches(indent_level_inches)
            pPr = p._p.get_or_add_pPr()
            pBdr, bottom = OxmlElement("w:pBdr"), OxmlElement("w:bottom")
            logger.debug("Adding <hr> as paragraph bottom border.")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # New explicit handling for known inline tags found at block
        # level
        elif tag_name in [
            "sup",
            "sub",
            "span",
            "font",
            "b",
            "i",
            "strong",
            "em",
            "u",
            "s",
            "a",
            "del",
            "strike",
            "br",
        ]:
            target_p_for_inline_tag = None
            is_first_run_in_target_p = True
            if hasattr(parent_docx_object, "_tc"):  # Parent is a cell
                if parent_docx_object.paragraphs:
                    target_p_for_inline_tag = parent_docx_object.paragraphs[-1]
                    is_first_run_in_target_p = not bool(
                        target_p_for_inline_tag.runs
                    )
                else:
                    target_p_for_inline_tag = parent_docx_object.add_paragraph()
            elif hasattr(
                parent_docx_object, "add_run"
            ):  # Parent is already a paragraph
                target_p_for_inline_tag = parent_docx_object
                is_first_run_in_target_p = not bool(
                    target_p_for_inline_tag.runs
                )
            else:  # Fallback
                target_p_for_inline_tag = self._create_paragraph_for_block(
                    parent_docx_object, indent_level_inches
                )

            self._process_inline_content(
                element,  # The inline Tag itself (e.g., <sup>, <br>)
                target_p_for_inline_tag,
                active_format_tags,  # Pass current active formatting context
                first=is_first_run_in_target_p,
                crt_par_classes=crt_par_classes,  # Pass through if available
            )

        else:
            # Default handling for unrecognized block tags or block-ish wrappers
            # (e.g. a div not handled above, or custom tags treated as blocks)
            new_active_tags = list(active_format_tags)
            if element.name.lower() in ["span", "font", "div"]:
                new_active_tags.append(element)

            # For children, decide if they are block (recursive call) or inline
            # (needs paragraph)
            created_para_for_inline = None
            for i, child in enumerate(element.children):
                if isinstance(child, Tag) and child.name.lower() in [
                    "p",
                    "div",
                    "h1",
                    "h2",
                    "h3",
                    "h4",
                    "h5",
                    "h6",
                    "ul",
                    "ol",
                    "table",
                    "img",
                    "svg",
                    "hr",
                    "dl",
                ]:
                    self._process_block_element(
                        child,
                        parent_docx_object,
                        new_active_tags,
                        indent_level_inches=indent_level_inches,
                        crt_par_classes=crt_par_classes,
                    )
                    created_para_for_inline = (
                        None  # Reset: next inline content needs new para
                    )
                elif isinstance(child, (NavigableString, Tag)):
                    # Inline or text found at block level, needs a paragraph
                    # context.
                    if created_para_for_inline is None:
                        # If parent_docx_object is a cell with existing
                        # paragraphs, and the current element (wrapper like
                        # <sup>) is inline, use the cell's last paragraph.
                        is_inline_wrapper = element.name.lower() in [
                            "sup",
                            "sub",
                            "span",
                            "font",
                            "b",
                            "i",
                            "strong",
                            "em",
                            "u",
                            "s",
                            "a",
                            "del",
                            "strike",
                        ]
                        if (
                            hasattr(parent_docx_object, "_tc")
                            and parent_docx_object.paragraphs
                            and is_inline_wrapper
                        ):
                            created_para_for_inline = (
                                parent_docx_object.paragraphs[-1]
                            )
                        else:
                            # Element is some other unknown block-ish wrapper,
                            # or parent is not a cell / has no paras / element
                            # is not inline wrapper.
                            created_para_for_inline = (
                                self._create_paragraph_for_block(
                                    parent_docx_object, indent_level_inches
                                )
                            )

                    self._process_inline_content(
                        child,
                        created_para_for_inline,
                        new_active_tags,
                        first=(i == 0)
                        and (
                            not bool(created_para_for_inline.runs)
                        ),  # Check if para is empty
                        is_dt_content=False,  # Assuming default context here
                        crt_par_classes=crt_par_classes,
                    )

    def _create_paragraph_for_block(
        self, parent_docx_object: Any, indent_level_inches: float | None
    ):
        """Helper to create a paragraph, typically for a new block element,
        applying indentation.
        """
        para = None
        if hasattr(parent_docx_object, "add_paragraph") and not isinstance(
            parent_docx_object, DocumentObject
        ):
            para = parent_docx_object.add_paragraph()  # Usually a Cell
            logger.debug(
                "Created paragraph in cell/block object for new block element."
            )
        else:
            # Document object, or parent is already a paragraph (less common
            # for new block)
            para = self.doc.add_paragraph()
            logger.debug(
                "Created paragraph in document root for new block element."
            )

        if indent_level_inches is not None and indent_level_inches > 0:
            para.paragraph_format.left_indent = Inches(indent_level_inches)
            logger.debug(
                "Applied indent of %s inches to new block paragraph.",
                indent_level_inches,
            )
        return para

    def _handle_dl_element(
        self,
        dl_element: Tag,
        parent_docx_object: Any,
        active_format_tags: list[Tag],
        current_indent_inches: float | None,
        crt_par_classes: list[str] | None = None,
    ):
        """Handles <dl> elements: styled <dt> and indented <dd> blocks."""
        logger.debug(
            "Processing <dl> element. Indent: %s", current_indent_inches
        )
        for child_node in dl_element.children:
            if not isinstance(child_node, Tag):
                continue

            child_tag_name = child_node.name.lower()

            if child_tag_name == "dt":
                logger.debug("  Processing <dt> in <dl>")
                p_dt = self._create_paragraph_for_block(
                    parent_docx_object, current_indent_inches
                )
                for i, dt_content_child in enumerate(child_node.children):
                    self._process_inline_content(
                        dt_content_child,
                        p_dt,
                        active_format_tags,
                        first=(i == 0),
                        is_dt_content=True,
                        crt_par_classes=crt_par_classes,
                    )

            elif child_tag_name == "dd":  # child_node is the <dd> Tag
                dd_children_indent_val = (current_indent_inches or 0) + 0.25
                logger.debug(
                    "  Processing <dd> in <dl>. New indent: %s",
                    dd_children_indent_val,
                )

                if not list(child_node.children):  # If DD is empty, skip
                    logger.debug("    <dd> is empty, skipping.")
                    continue

                # Create the primary paragraph for this <dd>'s inline content
                # stream
                current_dd_paragraph = self._create_paragraph_for_block(
                    parent_docx_object, dd_children_indent_val
                )
                is_first_run_in_dd_para = True

                for (
                    dd_content_item
                ) in child_node.children:  # Iterate children of <dd>
                    if isinstance(dd_content_item, (NavigableString, Tag)):
                        is_item_a_block_within_dd = False
                        if isinstance(dd_content_item, Tag):
                            item_tag_name = dd_content_item.name.lower()
                            # Define tags that should be treated as block-level
                            # when direct children of DD
                            if item_tag_name in [
                                "p",
                                "div",
                                "h1",
                                "h2",
                                "h3",
                                "h4",
                                "h5",
                                "h6",
                                "ul",
                                "ol",
                                "dl",
                                "table",
                                "hr",
                            ]:
                                logger.debug(
                                    "    <dd> child is block-level: <%s>",
                                    item_tag_name,
                                )
                                is_item_a_block_within_dd = True

                        if is_item_a_block_within_dd:
                            # Finalize current_dd_paragraph or remove if empty
                            if current_dd_paragraph:
                                if (
                                    not current_dd_paragraph.text.strip()
                                    and not current_dd_paragraph.runs
                                ):
                                    p_elem = current_dd_paragraph._element
                                    if p_elem.getparent() is not None:
                                        p_elem.getparent().remove(p_elem)
                                        logger.debug(
                                            "    Removed empty paragraph from "
                                            "<dd> before processing block "
                                            "child."
                                        )
                                # Set to None to signal a new one is needed if
                                # more inline content follows this block
                                current_dd_paragraph = None

                            # Process this block element. Its parent is the
                            # DL's parent. It will create its own paragraphs,
                            # using dd_children_indent_val.
                            self._process_block_element(
                                dd_content_item,  # The block Tag
                                parent_docx_object,
                                active_format_tags,
                                indent_level_inches=dd_children_indent_val,
                                crt_par_classes=crt_par_classes,
                            )
                            # After a block, the next inline item will need a
                            # new paragraph.
                            is_first_run_in_dd_para = (
                                True  # Reset for the potentially new paragraph
                            )
                            logger.debug(
                                "    Finished processing block child in <dd>. "
                                "Next inline needs new para."
                            )
                        else:
                            # This dd_content_item is inline (NavigableString
                            # or inline Tag like <sup>)
                            # It should go into the current_dd_paragraph.
                            logger.debug(
                                "    <dd> child is inline: %s",
                                str(dd_content_item)[:50].replace("\n", " "),
                            )
                            if (
                                current_dd_paragraph is None
                            ):  # Means a block was just processed
                                current_dd_paragraph = (
                                    self._create_paragraph_for_block(
                                        parent_docx_object,
                                        dd_children_indent_val,
                                    )
                                )
                                is_first_run_in_dd_para = True

                            self._process_inline_content(
                                # The NavigableString or inline Tag
                                dd_content_item,
                                current_dd_paragraph,
                                active_format_tags,
                                first=is_first_run_in_dd_para,
                                crt_par_classes=crt_par_classes,
                            )
                            # Update is_first_run_in_dd_para based on whether
                            # content was added
                            if (
                                isinstance(dd_content_item, NavigableString)
                                and str(dd_content_item).strip()
                            ):
                                is_first_run_in_dd_para = False
                            elif isinstance(dd_content_item, Tag):
                                # Consider content-ful tags or explicit line
                                # breaks as "content added"
                                if dd_content_item.get_text(
                                    strip=True
                                ) or dd_content_item.name.lower() in [
                                    "br",
                                    "img",
                                    "svg",
                                ]:
                                    is_first_run_in_dd_para = False
                                    logger.debug(
                                        "    Inline content added to <dd> "
                                        "paragraph, is_first_run_in_dd_para "
                                        "set to False."
                                    )
                    # else: ignore other PageElement types like Comment, etc.

                # After processing all children of <dd>, if the last
                # current_dd_paragraph is empty, remove it.
                if (
                    current_dd_paragraph
                    and not current_dd_paragraph.text.strip()
                    and not current_dd_paragraph.runs
                ):
                    p_elem = current_dd_paragraph._element
                    if (
                        p_elem.getparent() is not None
                    ):  # Ensure it's still part of the tree
                        p_elem.getparent().remove(p_elem)
                        logger.debug(
                            "  Removed last empty paragraph from <dd> after "
                            "processing all its children."
                        )
        logger.debug("Finished processing <dl> element.")

    def _get_attribute_as_int(
        self, element: Tag, attr_name: str, default_val: int = 1
    ) -> int:
        """Helper to safely get an attribute value as an integer."""
        attr_val = element.get(attr_name)
        if attr_val is None:
            return default_val

        val_to_convert = attr_val
        if isinstance(attr_val, list):
            if not attr_val:  # Empty list
                return default_val
            val_to_convert = str(attr_val[0])  # Take first item if list
        else:
            val_to_convert = str(attr_val)

        try:
            return int(val_to_convert)
        except ValueError:
            logger.warning(
                f"Invalid value for attribute {attr_name}: {val_to_convert}. "
                f"Using default {default_val}."
            )
            return default_val
