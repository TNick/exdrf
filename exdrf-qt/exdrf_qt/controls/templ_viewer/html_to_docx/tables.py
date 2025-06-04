import logging
from typing import TYPE_CHECKING, Any, List, Union, cast

from attrs import define, field
from bs4 import Tag
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import RGBColor
from docx.table import Table, _Cell
from lxml.etree import _Element  # type: ignore

if TYPE_CHECKING:
    from exdrf_qt.controls.templ_viewer.html_to_docx.main import (
        HtmlToDocxConverter,
    )

logger = logging.getLogger(__name__)
TABLE_STRIPED_BG_COLOR_RGB = (0xF2, 0xF2, 0xF2)  # Light gray for striping
OxmlElementType = Union[BaseOxmlElement, _Element]
CLASS_COLOR_MAP = {
    # "table-light": None,
    "table-success": "#d1e7dd",
    "bg-success": "#d1e7dd",
    "table-primary": "#cfe2ff",
    "bg-primary": "#cfe2ff",
    "table-danger": "#f8d7da",
    "bg-danger": "#f8d7da",
    "table-warning": "#fff3cd",
    "bg-warning": "#fff3cd",
    "table-info": "#cff4fc",
    "bg-info": "#cff4fc",
    "table-light": "#f8f9fa",
    "bg-light": "#f8f9fa",
    "table-secondary": "#e2e3e5",
    "bg-secondary": "#e2e3e5",
}


@define
class TableData:
    """Keeps track of the table data as we process the HTML.

    Attributes:
        html_grid: The grid for the docx table.
        html_rows: The rows from the HTML table.
        max_cols: The maximum number of columns across all rows.
        table_styles: The styles for the table.
        table_classes: The classes for the table.
        grid_r_idx: The current row index in the doc table.
        doc_table: The docx table object.
        num_logical_rows: The number of logical rows in the table.
    """

    cv: "HtmlToDocxConverter"
    html_grid: List[List[Tag | str | None]] = field(factory=list)
    html_rows: List[Tag] = field(factory=list)
    max_cols: int = field(default=0)
    table_styles: dict[str, Any] = field(factory=dict)
    table_classes: list[str] = field(factory=list)
    grid_r_idx: int = -1
    doc_table: Table = cast(Table, None)
    num_logical_rows: int = 0

    def _collect_table_rows(self, table_element: Tag):
        """Collect all rows from a table element."""

        # Collect rows from the table header.
        thead = table_element.find("thead", recursive=False)
        if thead and isinstance(thead, Tag):
            for tag in thead.find_all("tr", recursive=False):
                if isinstance(tag, Tag):
                    self.html_rows.append(tag)

        # Collect rows from the table body.
        tbody_elements = table_element.find_all("tbody", recursive=False)
        if tbody_elements:
            for tbody in tbody_elements:
                if isinstance(tbody, Tag):
                    for tag in tbody.find_all("tr", recursive=False):
                        if isinstance(tag, Tag):
                            self.html_rows.append(tag)

        # Collect rows from within the table directly.
        for tag in table_element.find_all("tr", recursive=False):
            if isinstance(tag, Tag):
                self.html_rows.append(tag)

        # Collect rows from the table footer.
        tfoot = table_element.find("tfoot", recursive=False)
        if tfoot and isinstance(tfoot, Tag):
            for tag in tfoot.find_all("tr", recursive=False):
                if isinstance(tag, Tag):
                    self.html_rows.append(tag)

        logger.debug(
            "Collected %d HTML rows for the table.", len(self.html_rows)
        )

    def _handle_table_row(self, r_idx, tr_element_maybe_str) -> None:
        if not isinstance(tr_element_maybe_str, Tag):
            logger.debug(
                "Skipping non-<tr> element at index %d: %s",
                r_idx,
                str(tr_element_maybe_str)[:50],
            )
            return

        tr_element: Tag = tr_element_maybe_str

        # This is the ID that we created ourselves in javascript.
        raw_tr_id = tr_element.get("data-docgen-id")
        tr_id_str: str | None = None
        if isinstance(raw_tr_id, list):
            if raw_tr_id:
                tr_id_str = str(raw_tr_id[0])
        elif isinstance(raw_tr_id, str):
            tr_id_str = raw_tr_id
        else:
            logger.warning(
                "Unexpected type for tr_id '%s': %s",
                raw_tr_id,
                type(raw_tr_id),
            )

        # Skip this entire row if the row is hidden.
        if tr_id_str:
            tr_element_data = self.cv.elements_map.get(tr_id_str)
            if tr_element_data and tr_element_data.get("is_hidden"):
                logger.debug(
                    "Skipping hidden <tr> element with ID %s", tr_id_str
                )
                return

        # Increment grid_r_idx and append row only for visible <tr>
        self.grid_r_idx += 1
        while len(self.html_grid) < self.grid_r_idx + 1:
            self.html_grid.append([])
        current_col_idx = 0

        # Process each cell in the row.
        for td_th_element in tr_element.find_all(["td", "th"], recursive=False):
            if not isinstance(td_th_element, Tag):
                logger.warning(
                    "Skipping non-<td>/<th> element at index %d: %s",
                    r_idx,
                    str(td_th_element)[:50],
                )
                return

            # Move to the next available column in the grid
            while (
                len(self.html_grid[self.grid_r_idx]) > current_col_idx
                and self.html_grid[self.grid_r_idx][current_col_idx] is not None
            ):
                current_col_idx += 1
            logger.debug(
                "  Cell processing: grid_r_idx=%d, current_col_idx=%d",
                self.grid_r_idx,
                current_col_idx,
            )

            # Get colspan and rowspan attributes
            colspan = self.cv._get_attribute_as_int(td_th_element, "colspan", 1)
            rowspan = self.cv._get_attribute_as_int(td_th_element, "rowspan", 1)
            logger.debug(
                "    HTML Cell (%s): colspan=%d, rowspan=%d",
                str(td_th_element.name),
                colspan,
                rowspan,
            )

            # Process rowspan
            for i in range(rowspan):
                target_r_in_grid = self.grid_r_idx + i

                # Ensure rows up to the target row exist
                self.html_grid.extend(
                    [
                        []
                        for _ in range(
                            len(self.html_grid), target_r_in_grid + 1
                        )
                    ]
                )

                # Ensure enough columns in the target row
                while (
                    len(self.html_grid[target_r_in_grid])
                    < current_col_idx + colspan
                ):
                    self.html_grid[target_r_in_grid].append(None)

                # Assign cell or mark as SPAN
                for j in range(colspan):
                    cell_value = td_th_element if i == 0 and j == 0 else "SPAN"
                    self.html_grid[target_r_in_grid][
                        current_col_idx + j
                    ] = cell_value

            # Update column index for the next cell
            current_col_idx += colspan

            # Track maximum number of columns across all rows.
            self.max_cols = max(self.max_cols, current_col_idx)

    def _merged_cell(
        self,
        doc_cell: _Cell,
        r_idx_grid_local: int,
        c_idx: int,
        colspan: int,
        rowspan: int,
    ) -> None:
        end_r_idx = min(
            r_idx_grid_local + rowspan - 1, self.num_logical_rows - 1
        )
        end_c_idx = min(c_idx + colspan - 1, self.max_cols - 1)
        if end_r_idx > r_idx_grid_local or end_c_idx > c_idx:
            try:
                doc_cell.merge(self.doc_table.cell(end_r_idx, end_c_idx))
                logger.debug(
                    "Merged cell at (grid_row=%d, col=%d) to "
                    "(grid_row=%d, col=%d)",
                    r_idx_grid_local,
                    c_idx,
                    end_r_idx,
                    end_c_idx,
                )
            except Exception as e:
                logger.warning(f"Cell merge failed: {e}")

    def _handle_cell(self, r_idx_grid_local: int, c_idx) -> None:
        html_cell_content = self.html_grid[r_idx_grid_local][c_idx]
        if html_cell_content is None or html_cell_content == "SPAN":
            return

        if not isinstance(html_cell_content, Tag):
            return

        html_cell_element: Tag = html_cell_content
        logger.debug(
            "Processing doc_cell at (grid_row=%d, col=%d) for " "HTML cell: %s",
            r_idx_grid_local,
            c_idx,
            str(html_cell_element)[:50],
        )
        doc_cell = self.doc_table.cell(r_idx_grid_local, c_idx)

        colspan = self.cv._get_attribute_as_int(html_cell_element, "colspan", 1)
        rowspan = self.cv._get_attribute_as_int(html_cell_element, "rowspan", 1)

        if rowspan > 1 or colspan > 1:
            self._merged_cell(
                doc_cell,
                r_idx_grid_local,
                c_idx,
                colspan,
                rowspan,
            )

        # Only remove the default <w:p> if we actually have content to add
        # i.e. if the HTML <td>/<th> has at least one child (text or tag)
        if html_cell_element is not None and list(html_cell_element.children):
            if doc_cell.paragraphs and doc_cell.paragraphs[0].text == "":
                p_element = doc_cell.paragraphs[0]._element
                p_element.getparent().remove(p_element)

        cell_styles, cell_classes = self.cv._parse_styles(html_cell_element)

        # Table striping and background color
        is_striped_table = "table-striped" in self.table_classes
        if is_striped_table and (r_idx_grid_local % 2 != 0):
            if not cell_styles.get("background-color"):
                shade_color = RGBColor(*TABLE_STRIPED_BG_COLOR_RGB)
                self.cv._set_cell_shading(doc_cell, str(shade_color))
        bg_color_str = cell_styles.get(
            "background-color"
        ) or self.table_styles.get("background-color")
        if not bg_color_str:
            for cls in self.table_classes:
                bg_color_str = CLASS_COLOR_MAP.get(cls)
                if bg_color_str:
                    break

        if bg_color_str:
            self.cv._set_cell_shading(doc_cell, bg_color_str)

        # Populate the cell with content from html_cell_element
        if html_cell_element and hasattr(html_cell_element, "children"):
            active_tags_for_cell_content = [html_cell_element]
            for child_node in html_cell_element.children:
                self.cv._process_block_element(
                    child_node, doc_cell, active_tags_for_cell_content
                )

    def set_table_borders(self):
        """
        Adds single-line borders to all cells in the given table.
        """
        from exdrf_qt.controls.templ_viewer.html_to_docx.main import (
            BORDER_COLOR_SUCCESS_RGB,
            DEFAULT_BORDER_COLOR_RGB,
            DEFAULT_BORDER_WIDTH_PT,
        )

        tbl = self.doc_table._tbl
        base_r, base_g, base_b = DEFAULT_BORDER_COLOR_RGB
        if "border-success" in self.table_classes:
            base_r, base_g, base_b = BORDER_COLOR_SUCCESS_RGB

        border_opacity = self.table_styles.get("border_opacity", 1.0)
        if not isinstance(border_opacity, (float, int)) or not (
            0.0 <= border_opacity <= 1.0
        ):
            border_opacity = 1.0
        if border_opacity < 0.05:  # Effectively transparent, skip
            return

        parsed_width_val = DEFAULT_BORDER_WIDTH_PT
        border_sz = max(1, int(parsed_width_val * 8))

        # Look for an existing <w:tblPr> element; if missing,
        # create and insert it.
        tblPr = tbl.find(qn("w:tblPr"))  # type: ignore
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Create <w:tblBorders> element
        tblBorders = OxmlElement("w:tblBorders")

        # Define border attributes for each side
        for border_name in (
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ):
            border_el = OxmlElement(f"w:{border_name}")

            self._set_cell_border_color(
                border_el,
                (base_r, base_g, base_b),
                size_pt=border_sz,
                alpha=border_opacity,
            )

            tblBorders.append(border_el)

        tblPr.append(tblBorders)

    def _set_cell_border_color(
        self,
        border_side_element: OxmlElementType,
        color_rgb: tuple[int, int, int],
        size_pt: int = 4,
        alpha: float = 1.0,
    ):
        """Set the borders of the cell.

        Args:
            border_side_element: The border side element to set the color of.
            color_rgb: The color of the border in RGB format.
            size_pt: The size of the border in points. This is w:sz unit
                (eighths of a point).
            alpha: The opacity of the border (0.0 to 1.0).
        """
        r, g, b = color_rgb
        actual_r, actual_g, actual_b = r, g, b

        # Apply opacity by blending with white
        if alpha < 1.0 and alpha >= 0.0:
            actual_r = int(r * alpha + 255 * (1 - alpha))
            actual_g = int(g * alpha + 255 * (1 - alpha))
            actual_b = int(b * alpha + 255 * (1 - alpha))

        # Ensure values are within valid range
        final_r = max(0, min(actual_r, 255))
        final_g = max(0, min(actual_g, 255))
        final_b = max(0, min(actual_b, 255))

        # Convert to hex format expected by Word XML (without # prefix)
        hex_color = f"{final_r:02X}{final_g:02X}{final_b:02X}"

        border_side_element.set(qn("w:val"), "single")
        border_side_element.set(qn("w:sz"), str(size_pt))
        border_side_element.set(qn("w:color"), hex_color)
        border_side_element.set(qn("w:space"), "0")

    def set_full_width(self):
        """Set the table to be the full width of the page."""
        tbl = self.doc_table._tbl

        # Look for an existing <w:tblPr> element; if missing,
        # create and insert it.
        tblPr = tbl.find(qn("w:tblPr"))  # type: ignore
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblW = tblPr.find(qn("w:tblW"))  # type: ignore
        if tblW is not None:
            tblPr.remove(tblW)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "5000")  # 5000 twentieths of a percent = 100%
        tblW.set(qn("w:type"), "pct")  # pct means percentage-based width
        tblPr.append(tblW)

        # Disable automatic resizing so Word respects the 100% width
        self.doc_table.autofit = False

    def process(self, table_element: Tag, parent_docx_object):
        """Process a table element and convert it to a Docx table."""
        logger.debug(
            "--- Starting _handle_table for: %s ---", str(table_element)[:100]
        )

        # Get table classes
        self.table_styles, self.table_classes = self.cv._parse_styles(
            table_element
        )
        logger.debug(
            "Table styles: %s, Table classes: %s",
            self.table_styles,
            self.table_classes,
        )

        # Collect html rows.
        self._collect_table_rows(table_element)

        # Process each row.
        for r_idx, tr_element_maybe_str in enumerate(self.html_rows):
            self._handle_table_row(r_idx, tr_element_maybe_str)

        # Ensure the grid has enough columns in every row.
        self.num_logical_rows = len(self.html_grid)
        for r_list in self.html_grid:
            while len(r_list) < self.max_cols:
                r_list.append(None)

        # If the table has 0 rows or 0 columns, skip it.
        if self.num_logical_rows == 0 or self.max_cols == 0:
            logger.warning(
                "Table has 0 rows or 0 columns after processing grid. "
                "Skipping table."
            )
            return

        self.doc_table = parent_docx_object.add_table(
            rows=self.num_logical_rows,
            cols=self.max_cols,
        )
        # Apply table-level styles like table-layout: fixed if needed (not
        # requested yet) doc_table.autofit = False doc_table.layout_type =
        # WD_TABLE_LAYOUT.FIXED

        for r_idx_grid_local in range(self.num_logical_rows):
            for c_idx in range(self.max_cols):
                self._handle_cell(r_idx_grid_local, c_idx)
        self.set_table_borders()
        self.set_full_width()

        logger.debug(
            "--- Finished _handle_table for: %s ---", str(table_element)[:100]
        )
