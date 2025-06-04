import io
import logging
import re
from typing import Any, Optional, Union

from bs4 import BeautifulSoup, FeatureNotFound, Tag
from bs4.element import NavigableString, PageElement
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from minify_html import minify
from PyQt5.QtCore import QBuffer, QByteArray, QIODevice, QUrl
from PyQt5.QtGui import QDesktopServices, QImage
from PyQt5.QtWebEngineWidgets import QWebEngineView

from exdrf_qt.controls.templ_viewer.html_to_docx.screen_grabber import (
    FullPageGrabber,
)
from exdrf_qt.controls.templ_viewer.html_to_docx.tables import TableData

logger = logging.getLogger(__name__)

# Define constants for styling
# Store colors as (R, G, B) tuples for opacity processing
# And create RGBColor objects on demand.
BORDER_COLOR_SUCCESS_RGB = (0x19, 0x87, 0x54)  # Bootstrap's success green
DEFAULT_BORDER_COLOR_RGB = (0x00, 0x00, 0x00)  # Black
DEFAULT_BORDER_WIDTH_PT = 0.75  # Approx 1px, use in Pt() later

default_styles = """
No List (LIST (4))

Header (PARAGRAPH (1))
Footer (PARAGRAPH (1))

Normal (PARAGRAPH (1))
No Spacing (PARAGRAPH (1))

Title (PARAGRAPH (1))
Subtitle (PARAGRAPH (1))
Heading 1 (PARAGRAPH (1))
Heading 2 (PARAGRAPH (1))
Heading 3 (PARAGRAPH (1))
Heading 4 (PARAGRAPH (1))
Heading 5 (PARAGRAPH (1))
Heading 6 (PARAGRAPH (1))
Heading 7 (PARAGRAPH (1))
Heading 8 (PARAGRAPH (1))
Heading 9 (PARAGRAPH (1))

Body Text (PARAGRAPH (1))
Body Text 2 (PARAGRAPH (1))
Body Text 3 (PARAGRAPH (1))

List Paragraph (PARAGRAPH (1))
List (PARAGRAPH (1))
List 2 (PARAGRAPH (1))
List 3 (PARAGRAPH (1))
List Bullet (PARAGRAPH (1))
List Bullet 2 (PARAGRAPH (1))
List Bullet 3 (PARAGRAPH (1))
List Number (PARAGRAPH (1))
List Number 2 (PARAGRAPH (1))
List Number 3 (PARAGRAPH (1))
List Continue (PARAGRAPH (1))
List Continue 2 (PARAGRAPH (1))
List Continue 3 (PARAGRAPH (1))

macro (PARAGRAPH (1))
Quote (PARAGRAPH (1))
Intense Quote (PARAGRAPH (1))
Caption (PARAGRAPH (1))

TOC Heading (PARAGRAPH (1))

------------ CHARACTERS ------------

Header Char (CHARACTER (2))
Footer Char (CHARACTER (2))
Default Paragraph Font (CHARACTER (2))
Heading 1 Char (CHARACTER (2))
Heading 2 Char (CHARACTER (2))
Heading 3 Char (CHARACTER (2))
Title Char (CHARACTER (2))
Subtitle Char (CHARACTER (2))
Body Text Char (CHARACTER (2))
Body Text 2 Char (CHARACTER (2))
Body Text 3 Char (CHARACTER (2))
Macro Text Char (CHARACTER (2))
Quote Char (CHARACTER (2))
Heading 4 Char (CHARACTER (2))
Heading 5 Char (CHARACTER (2))
Heading 6 Char (CHARACTER (2))
Heading 7 Char (CHARACTER (2))
Heading 8 Char (CHARACTER (2))
Heading 9 Char (CHARACTER (2))
Strong (CHARACTER (2))
Emphasis (CHARACTER (2))
Intense Quote Char (CHARACTER (2))
Subtle Emphasis (CHARACTER (2))
Intense Emphasis (CHARACTER (2))
Subtle Reference (CHARACTER (2))
Intense Reference (CHARACTER (2))
Book Title (CHARACTER (2))

------------ TABLES ------------

Normal Table (TABLE (3))
Table Grid (TABLE (3))
Light Shading (TABLE (3))
Light Shading Accent 1 (TABLE (3))
Light Shading Accent 2 (TABLE (3))
Light Shading Accent 3 (TABLE (3))
Light Shading Accent 4 (TABLE (3))
Light Shading Accent 5 (TABLE (3))
Light Shading Accent 6 (TABLE (3))
Light List (TABLE (3))
Light List Accent 1 (TABLE (3))
Light List Accent 2 (TABLE (3))
Light List Accent 3 (TABLE (3))
Light List Accent 4 (TABLE (3))
Light List Accent 5 (TABLE (3))
Light List Accent 6 (TABLE (3))
Light Grid (TABLE (3))
Light Grid Accent 1 (TABLE (3))
Light Grid Accent 2 (TABLE (3))
Light Grid Accent 3 (TABLE (3))
Light Grid Accent 4 (TABLE (3))
Light Grid Accent 5 (TABLE (3))
Light Grid Accent 6 (TABLE (3))
Medium Shading 1 (TABLE (3))
Medium Shading 1 Accent 1 (TABLE (3))
Medium Shading 1 Accent 2 (TABLE (3))
Medium Shading 1 Accent 3 (TABLE (3))
Medium Shading 1 Accent 4 (TABLE (3))
Medium Shading 1 Accent 5 (TABLE (3))
Medium Shading 1 Accent 6 (TABLE (3))
Medium Shading 2 (TABLE (3))
Medium Shading 2 Accent 1 (TABLE (3))
Medium Shading 2 Accent 2 (TABLE (3))
Medium Shading 2 Accent 3 (TABLE (3))
Medium Shading 2 Accent 4 (TABLE (3))
Medium Shading 2 Accent 5 (TABLE (3))
Medium Shading 2 Accent 6 (TABLE (3))
Medium List 1 (TABLE (3))
Medium List 1 Accent 1 (TABLE (3))
Medium List 1 Accent 2 (TABLE (3))
Medium List 1 Accent 3 (TABLE (3))
Medium List 1 Accent 4 (TABLE (3))
Medium List 1 Accent 5 (TABLE (3))
Medium List 1 Accent 6 (TABLE (3))
Medium List 2 (TABLE (3))
Medium List 2 Accent 1 (TABLE (3))
Medium List 2 Accent 2 (TABLE (3))
Medium List 2 Accent 3 (TABLE (3))
Medium List 2 Accent 4 (TABLE (3))
Medium List 2 Accent 5 (TABLE (3))
Medium List 2 Accent 6 (TABLE (3))
Medium Grid 1 (TABLE (3))
Medium Grid 1 Accent 1 (TABLE (3))
Medium Grid 1 Accent 2 (TABLE (3))
Medium Grid 1 Accent 3 (TABLE (3))
Medium Grid 1 Accent 4 (TABLE (3))
Medium Grid 1 Accent 5 (TABLE (3))
Medium Grid 1 Accent 6 (TABLE (3))
Medium Grid 2 (TABLE (3))
Medium Grid 2 Accent 1 (TABLE (3))
Medium Grid 2 Accent 2 (TABLE (3))
Medium Grid 2 Accent 3 (TABLE (3))
Medium Grid 2 Accent 4 (TABLE (3))
Medium Grid 2 Accent 5 (TABLE (3))
Medium Grid 2 Accent 6 (TABLE (3))
Medium Grid 3 (TABLE (3))
Medium Grid 3 Accent 1 (TABLE (3))
Medium Grid 3 Accent 2 (TABLE (3))
Medium Grid 3 Accent 3 (TABLE (3))
Medium Grid 3 Accent 4 (TABLE (3))
Medium Grid 3 Accent 5 (TABLE (3))
Medium Grid 3 Accent 6 (TABLE (3))
Dark List (TABLE (3))
Dark List Accent 1 (TABLE (3))
Dark List Accent 2 (TABLE (3))
Dark List Accent 3 (TABLE (3))
Dark List Accent 4 (TABLE (3))
Dark List Accent 5 (TABLE (3))
Dark List Accent 6 (TABLE (3))
Colorful Shading (TABLE (3))
Colorful Shading Accent 1 (TABLE (3))
Colorful Shading Accent 2 (TABLE (3))
Colorful Shading Accent 3 (TABLE (3))
Colorful Shading Accent 4 (TABLE (3))
Colorful Shading Accent 5 (TABLE (3))
Colorful Shading Accent 6 (TABLE (3))
Colorful List (TABLE (3))
Colorful List Accent 1 (TABLE (3))
Colorful List Accent 2 (TABLE (3))
Colorful List Accent 3 (TABLE (3))
Colorful List Accent 4 (TABLE (3))
Colorful List Accent 5 (TABLE (3))
Colorful List Accent 6 (TABLE (3))
Colorful Grid (TABLE (3))
Colorful Grid Accent 1 (TABLE (3))
Colorful Grid Accent 2 (TABLE (3))
Colorful Grid Accent 3 (TABLE (3))
Colorful Grid Accent 4 (TABLE (3))
Colorful Grid Accent 5 (TABLE (3))
Colorful Grid Accent 6 (TABLE (3))
"""


class HtmlToDocxConverter:
    """Converts HTML to DOCX."""

    output_path: str
    view: QWebEngineView
    doc: DocumentObject
    image_cache: dict[str, bytes]
    full_page_qimage: QImage | None
    elements_map: dict[str, dict[str, Any]]
    page_width: int = -1
    page_height: int = -1
    grabber: Optional[FullPageGrabber] = None

    def __init__(self, view: QWebEngineView):
        self.view = view
        self.doc = Document()
        self.image_cache: dict[str, bytes] = {}
        self.full_page_qimage: QImage | None = None
        self.elements_map: dict[str, dict[str, Any]] = {}

    def export_to_docx(self, output_path: str):
        if self.grabber is not None:
            if self.grabber.view is not None:
                self.grabber.view.deleteLater()
            self.grabber = None

        page = self.view.page()
        assert page is not None
        page.toHtml(self.got_html)
        self.output_path = output_path
        logger.debug("Starting DOCX export to: %s", output_path)

    def got_html(self, html: Union[str, None]):
        if not html:
            logger.error("No HTML content received.")
            return
        logger.debug("Got HTML: %s", html)
        self.grabber = FullPageGrabber(
            html=html,
            callback=self.got_page_from_grabber,  # type: ignore
        )

    def got_page_from_grabber(self, grabber: FullPageGrabber):
        """The callback from the FullPageGrabber."""
        if grabber.errors:
            logger.error("Errors from grabber: %s", grabber.errors)
            return
        assert grabber.js_result is not None
        assert grabber.width is not None
        assert grabber.height is not None
        assert grabber.pixmap is not None
        assert grabber.html is not None

        self.elements_map = {}
        for item in grabber.js_result["elementInfoList"]:
            self.elements_map[item["id"]] = {
                "tagName": item["tagName"],
                "is_hidden": item["isEffectivelyHidden"],
                # Will be None if not present
                "geometry": item.get("geometry"),
            }
        self.page_width = grabber.width
        self.page_height = grabber.height

        self.full_page_qimage = grabber.pixmap.toImage()
        if self.full_page_qimage is not None and grabber.debug_mode:
            self.full_page_qimage.save("full_page_qimage.png")
        assert self.full_page_qimage is not None
        if self.full_page_qimage.isNull():
            logger.error("Failed to grab full page screenshot.")
            self.full_page_qimage = None
        else:
            logger.info(
                "Full page screenshot captured: %sx%s",
                self.full_page_qimage.width(),
                self.full_page_qimage.height(),
            )

        page = self.view.page()
        assert page is not None

        self.got_current_html_with_ids(grabber.js_result["fullHTMLTree"])

    def got_current_html_with_ids(self, html_content: str):
        logger.debug(
            "Received HTML content. Starting DOCX document processing."
        )
        self.doc = Document()
        self.image_cache = {}  # Reset image cache for this run
        logger.debug("Minifying HTML content...")
        html_content = minify(
            html_content,
            allow_noncompliant_unquoted_attribute_values=False,
            allow_optimal_entities=False,
            allow_removing_spaces_between_attributes=True,
            keep_closing_tags=True,
            keep_comments=False,
            keep_html_and_head_opening_tags=True,
            keep_input_type_text_attr=False,
            keep_ssi_comments=False,
            minify_css=True,
            minify_doctype=False,
            minify_js=True,
            preserve_brace_template_syntax=False,
            preserve_chevron_percent_template_syntax=False,
            remove_bangs=False,
            remove_processing_instructions=False,
        )
        try:
            soup = BeautifulSoup(html_content, "lxml")
        except FeatureNotFound:
            soup = BeautifulSoup(html_content, "html.parser")
        except Exception as e:
            logger.error("BeautifulSoup parsing error: %s", e)
            return
        logger.debug("HTML parsing complete. Decomposing script tags.")

        # Strip all <script> tags
        for script_tag in soup.find_all("script"):
            logger.debug("Removing script tag: %s", str(script_tag)[:100])
            script_tag.decompose()

        body = soup.find("body")
        process_root = body if body else soup

        for element in process_root.children:  # type: ignore
            if isinstance(element, (Tag, NavigableString)):
                logger.debug("Processing root child: %s", str(element)[:100])
                self._process_block_element(element, self.doc, [])

        try:
            logger.debug(
                "Attempting to save DOCX file to: %s", self.output_path
            )
            self.doc.save(self.output_path)
            logger.info("DOCX file successfully saved to %s", self.output_path)
            QDesktopServices.openUrl(QUrl.fromLocalFile(self.output_path))
        except Exception as e:
            logger.error(
                "Error saving DOCX file to %s: %s", self.output_path, e
            )

    def _add_element_as_image(self, element_id: str, parent_docx_object):
        # Use self.elements_map to get geometry
        if not self.full_page_qimage:
            logger.error(
                "Missing full page screenshot for element %s", element_id
            )
            return
        logger.debug("Adding element %s as image.", element_id)

        element_data = self.elements_map.get(element_id)
        if not element_data:
            logger.error("No map data found for element %s", element_id)
            return

        geometry_data = element_data.get("geometry")
        if not geometry_data:
            logger.error(
                "Missing geometry for image/svg element %s", element_id
            )
            return

        if element_id in self.image_cache:
            img_bytes = self.image_cache[element_id]
        else:
            logger.debug("Element %s not in image cache, cropping.", element_id)
            # Use geometry_data directly for coordinates and dimensions
            x = int(geometry_data["docX"])
            y = int(geometry_data["docY"])
            width = int(geometry_data["docWidth"])
            height = int(geometry_data["docHeight"])

            if width <= 0 or height <= 0:
                logger.error(
                    "Invalid dimensions for element %s: w=%s, h=%s",
                    element_id,
                    width,
                    height,
                )
                return

            cropped_qimage = self.full_page_qimage.copy(x, y, width, height)
            if cropped_qimage.isNull():
                logger.error("Failed to crop image for element %s", element_id)
                return
            logger.debug("Successfully cropped image for %s.", element_id)

            byte_array = QByteArray()
            buffer = QBuffer(byte_array)
            buffer.open(QIODevice.OpenModeFlag.WriteOnly)
            cropped_qimage.save(buffer, "PNG")
            img_bytes = bytes(byte_array.data())
            self.image_cache[element_id] = img_bytes

        if img_bytes:
            img_stream = io.BytesIO(img_bytes)
            # Use geometry_data directly for width/height
            w_px = geometry_data["docWidth"]
            h_px = geometry_data["docHeight"]

            para_for_img = None
            if hasattr(parent_docx_object, "add_paragraph") and not isinstance(
                parent_docx_object, DocumentObject
            ):
                # If parent is a cell or similar that can add paragraphs
                para_for_img = parent_docx_object.add_paragraph()
            elif hasattr(
                parent_docx_object, "add_run"
            ):  # Parent is already a paragraph
                para_for_img = parent_docx_object
            else:  # Document object or fallback
                para_for_img = self.doc.add_paragraph()

            try:
                dpi = 96.0
                w_inch = Inches(w_px / dpi) if w_px > 0 else None
                h_inch = Inches(h_px / dpi) if h_px > 0 else None

                run_for_picture = para_for_img.add_run()  # type: ignore
                if w_inch and h_inch:
                    run_for_picture.add_picture(
                        img_stream, width=w_inch, height=h_inch
                    )
                elif w_inch:
                    run_for_picture.add_picture(img_stream, width=w_inch)
                elif h_inch:
                    run_for_picture.add_picture(img_stream, height=h_inch)
                else:  # Add with original pixel size
                    run_for_picture.add_picture(img_stream)
            except Exception as e:
                logger.error(
                    "Error adding cropped image %s to DOCX: %s",
                    element_id,
                    e,
                )
            logger.debug("Image %s added to DOCX.", element_id)

    def _hex_to_rgb(self, hex_color: str) -> RGBColor | None:
        """Color and Style Parsers."""
        hex_color = hex_color.lstrip("#")
        if len(hex_color) == 3:
            hex_color = "".join([c * 2 for c in hex_color])
        if len(hex_color) == 6:
            try:
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                return RGBColor(r, g, b)
            except ValueError:
                return None
        return None

    def _get_color(self, css_color_string: str) -> RGBColor | None:
        css_color_string = css_color_string.lower().strip()
        named_colors = {
            "black": RGBColor(0, 0, 0),
            "white": RGBColor(255, 255, 255),
            "red": RGBColor(255, 0, 0),
            "green": RGBColor(0, 128, 0),
            "blue": RGBColor(0, 0, 255),
            "yellow": RGBColor(255, 255, 0),
            "transparent": None,
        }
        if css_color_string in named_colors:
            return named_colors[css_color_string]
        if css_color_string.startswith("#"):
            return self._hex_to_rgb(css_color_string)
        elif css_color_string.startswith("rgb("):
            try:
                match = re.match(
                    r"rgb\\((\\d+),\\s*(\\d+),\\s*(\\d+)\\)", css_color_string
                )
                if match:
                    r, g, b = map(int, match.groups())
                    return RGBColor(r, g, b)
            except Exception:
                return None
        return None

    def _parse_styles(self, element: Tag) -> tuple[dict[str, Any], list[str]]:
        styles: dict[str, Any] = {}
        style_attr = element.get("style")
        if style_attr and isinstance(style_attr, str):
            for style_declaration in style_attr.split(";"):
                style_declaration = style_declaration.strip()
                if ":" in style_declaration:
                    prop, val = style_declaration.split(":", 1)
                    prop = prop.strip().lower()
                    val = val.strip()
                    # logger.debug(f"Parsed style from element
                    #           {element.name}: {prop}={val}")
                    if prop == "--bs-border-opacity":
                        try:
                            styles["border_opacity"] = float(val)
                        except ValueError:
                            logger.warning("Invalid opacity: %s", val)
                    else:
                        styles[prop] = val

        final_class_list: list[str] = []
        class_attr = element.get("class")
        if isinstance(class_attr, str):
            final_class_list = class_attr.split()
        elif isinstance(class_attr, list):
            final_class_list = [
                str(c) for c in class_attr if isinstance(c, str)
            ]
        # if final_class_list:
        #     logger.debug(f"Parsed classes from element
        #           {element.name}: {final_class_list}")
        return styles, final_class_list

    def _apply_formatting_to_run(self, run, element_tag: Tag):
        """Content Processors"""
        tag_name = element_tag.name.lower()
        styles, _ = self._parse_styles(element_tag)

        if tag_name in ["b", "strong"] or styles.get("font-weight") == "bold":
            run.bold = True
        if tag_name in ["i", "em"] or styles.get("font-style") == "italic":
            run.italic = True

        text_decoration = styles.get("text-decoration", "")
        if tag_name == "u" or "underline" in text_decoration:
            run.underline = True
        if (
            tag_name in ["s", "strike", "del"]
            or "line-through" in text_decoration
        ):
            run.strike = True

        color_str = styles.get("color")
        if color_str:
            color = self._get_color(color_str)
            if color:
                logger.debug(
                    "Applying font color %s to run for element %s",
                    color,
                    element_tag.name,
                )
                run.font.color.rgb = color

        font_family_str = styles.get("font-family")
        if font_family_str:
            primary_font = (
                font_family_str.split(",")[0]
                .strip()
                .replace("'", "")
                .replace('"', "")
            )
            if primary_font:
                run.font.name = primary_font

        font_size_str = styles.get("font-size")
        if font_size_str:
            try:
                val_str = re.sub(r"[^\\d\\.]", "", font_size_str)
                if val_str:
                    val = float(val_str)
                    if "pt" in font_size_str:
                        run.font.size = Pt(val)
                    elif "px" in font_size_str:
                        run.font.size = Pt(val * 0.75)
            except ValueError:
                pass

        if tag_name == "a":
            if not run.font.color.rgb:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                logger.debug("Applied default link color to run for <a> tag")
            if not run.underline:
                run.underline = WD_UNDERLINE.SINGLE
                logger.debug(
                    "Applied default link underline to run for <a> tag"
                )

    def _process_inline_content(
        self,
        html_node: PageElement,
        docx_parent_paragraph,
        active_format_tags: list[Tag],
        first: bool = False,
        is_dt_content: bool = False,
        crt_par_classes: list[str] | None = None,
    ):
        if isinstance(html_node, NavigableString):
            text = str(html_node)
            if len(text) > 0:
                prefix = " " if text[0].isspace() and not first else ""
                suffix = " " if text[-1].isspace() else ""
                text = prefix + text.strip() + suffix
                logger.debug(
                    "Processing NavigableString: '%s' (first=%s, is_dt=%s)",
                    text[:50].replace("\n", " "),
                    first,
                    is_dt_content,
                )

                run = docx_parent_paragraph.add_run(text)
                if is_dt_content:
                    run.bold = True
                for fmt_tag in active_format_tags:
                    self._apply_formatting_to_run(run, fmt_tag)

        elif isinstance(html_node, Tag):
            # Check if element should be skipped based on d-none
            # Safely get element_id for map lookup
            raw_element_id = html_node.get("data-docgen-id")
            element_id_str: str | None = None
            if isinstance(raw_element_id, list):
                if raw_element_id:  # Check if list is not empty
                    element_id_str = str(raw_element_id[0])
            elif isinstance(raw_element_id, str):
                element_id_str = raw_element_id

            if element_id_str:
                element_data = self.elements_map.get(element_id_str)
                if element_data and element_data.get("is_hidden"):
                    logger.debug(
                        "Skip hidden inline %s (%s)",
                        element_id_str,
                        html_node.name,
                    )
                    return  # Skip this hidden element

            tag_name = html_node.name.lower()
            logger.debug("Processing inline tag: <%s>", tag_name)

            # Ignore <button> tags and their content
            if tag_name == "button":
                logger.debug(
                    "Ignoring <button> tag and its content in block "
                    "context: %s",
                    str(html_node)[:50],
                )
                return

            if tag_name == "br":
                docx_parent_paragraph.add_run().add_break()
                return
            logger.debug("Added <br> to paragraph.")

            if tag_name == "sup":
                sup_text = (
                    html_node.get_text().strip()
                )  # Get all text within <sup>
                if sup_text:  # Only add run if there's text
                    logger.debug("Applying <sup> with text: '%s'", sup_text)
                    run = docx_parent_paragraph.add_run(sup_text)
                    run.font.superscript = True
                    # Apply any inherited formatting from parent tags
                    for fmt_tag in active_format_tags:
                        self._apply_formatting_to_run(run, fmt_tag)
                    # Apply any direct styling on the <sup> tag itself
                    self._apply_formatting_to_run(run, html_node)
                return  # Consumed <sup> tag and its content
            logger.debug("Finished processing <sup> tag.")

            if tag_name in ["img", "svg"]:
                # Safely get element_id for map lookup
                raw_img_id = html_node.get("data-docgen-id")
                img_id_str: str | None = None
                if isinstance(raw_img_id, list):
                    if raw_img_id:  # Check if list is not empty
                        img_id_str = str(raw_img_id[0])
                elif isinstance(raw_img_id, str):
                    img_id_str = raw_img_id

                if img_id_str:
                    element_map_data = self.elements_map.get(img_id_str, {})
                    if element_map_data.get("geometry"):
                        self._add_element_as_image(
                            img_id_str, docx_parent_paragraph
                        )
                    else:
                        logger.debug(
                            f"Skip inline img/svg {img_id_str}, no geometry."
                        )
                else:
                    logger.error(
                        "Skip inline <%s> no data-docgen-id: %.30s",
                        tag_name,
                        str(html_node),
                    )
                return  # Consumed img/svg (or skipped)
            logger.debug("Finished processing <img> or <svg> tag.")

            current_active_tags = list(active_format_tags)
            is_inline_formatting_provider = tag_name in [
                "b",
                "strong",
                "i",
                "em",
                "u",
                "s",
                "strike",
                "del",
                "span",
                "font",
                "a",
            ]
            if is_inline_formatting_provider:
                current_active_tags.append(html_node)

            if tag_name == "a":
                href = html_node.get("href")
                link_text = "".join(
                    map(str, html_node.find_all(string=True, recursive=True))
                ).strip()  # Ensure strings
                if href and link_text:
                    try:
                        h_run = docx_parent_paragraph.add_hyperlink(
                            href, link_text, is_external=True
                        )
                        for fmt_tag in current_active_tags:
                            self._apply_formatting_to_run(h_run, fmt_tag)
                    except Exception as e:
                        logger.error(
                            "Failed to add hyperlink for %s: %s. "
                            "Adding as text.",
                            href,
                            e,
                        )
                        for i, child in enumerate(html_node.children):
                            self._process_inline_content(
                                child,
                                docx_parent_paragraph,
                                current_active_tags,
                                first=(i == 0),
                                is_dt_content=False,
                                crt_par_classes=crt_par_classes,
                            )
                    return
                # Fallback for 'a' if no href/text or error
                for i, child in enumerate(html_node.children):
                    self._process_inline_content(
                        child,
                        docx_parent_paragraph,
                        current_active_tags,
                        first=(i == 0),
                        is_dt_content=False,
                        crt_par_classes=crt_par_classes,
                    )
            else:
                for i, child in enumerate(html_node.children):
                    self._process_inline_content(
                        child,
                        docx_parent_paragraph,
                        current_active_tags,
                        first=(i == 0),
                        is_dt_content=False,
                        crt_par_classes=crt_par_classes,
                    )

    def _set_cell_shading(self, cell, color_str: str):
        color = self._get_color(color_str)
        if color and color_str.lower() != "transparent":
            try:
                logger.debug("Applying cell shading %s", color_str)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), str(color))
                shd.set(qn("w:val"), "clear")
                tcPr = cell._tc.get_or_add_tcPr()
                for existing_shd in tcPr.xpath("./w:shd"):
                    tcPr.remove(existing_shd)
                tcPr.append(shd)
            except Exception as e:
                logger.error(
                    "Error applying cell shading %s: %s",
                    color_str,
                    e,
                )

    def _handle_table(self, table_element: Tag, parent_docx_object):
        """Process a table element and convert it to a Docx table."""
        table_data = TableData(cv=self)
        table_data.process(table_element, parent_docx_object)

    def _process_block_element(
        self,
        element: PageElement,
        parent_docx_object,
        active_format_tags: list[Tag],
        indent_level_inches: float | None = None,
        crt_par_classes: list[str] | None = None,
    ):
        current_paragraph: Any = None

        if isinstance(element, NavigableString):
            # Delegate NavigableString to _process_inline_content
            # Ensure a paragraph context exists in parent_docx_object
            target_p_for_nav_str = None
            is_first_run = True
            logger.debug(
                "Processing NavigableString (block context): '%s'",
                str(element)[:50].replace("\n", " "),
            )
            if hasattr(parent_docx_object, "_tc"):  # Parent is a cell
                if parent_docx_object.paragraphs:
                    target_p_for_nav_str = parent_docx_object.paragraphs[-1]
                    is_first_run = not bool(target_p_for_nav_str.runs)
                else:
                    target_p_for_nav_str = parent_docx_object.add_paragraph()
            elif hasattr(
                parent_docx_object, "add_run"
            ):  # Parent is already a paragraph
                target_p_for_nav_str = parent_docx_object
                is_first_run = not bool(target_p_for_nav_str.runs)
            else:  # Fallback (e.g., document root)
                target_p_for_nav_str = self._create_paragraph_for_block(
                    parent_docx_object, indent_level_inches
                )

            self._process_inline_content(
                element,  # The NavigableString itself
                target_p_for_nav_str,
                active_format_tags,
                first=is_first_run,  # Pass 'first' status
                # is_dt_content and crt_par_classes ar e
                # context-dependent
            )
            return

        if not isinstance(element, Tag):
            # Ignore other PageElement types like Comment
            return

        # Check if element should be skipped based on d-none
        # Safely get element_id for map lookup
        raw_element_id_block = element.get("data-docgen-id")
        element_id_block_str: str | None = None
        if isinstance(raw_element_id_block, list):
            if raw_element_id_block:  # Check if list is not empty
                element_id_block_str = str(raw_element_id_block[0])
        elif isinstance(raw_element_id_block, str):
            element_id_block_str = raw_element_id_block

        if element_id_block_str:
            element_data = self.elements_map.get(element_id_block_str)
            if element_data and element_data.get("is_hidden"):
                logger.debug(
                    f"Skip hidden block {element_id_block_str} ({element.name})"
                )
                return  # Skip this hidden element

        tag_name = element.name.lower()
        logger.debug(
            "Processing block tag: <%s> with indent %s",
            tag_name,
            indent_level_inches,
        )

        # Ignore <button> tags and their content
        if tag_name == "button":
            if "include-btn-in-print" not in element.attrs["class"]:
                logger.debug(
                    "Ignoring <button> tag and its content in block "
                    "context: %s; can include by setting "
                    "include-btn-in-print class",
                    str(element)[:50],
                )
                return

        if tag_name in ["img", "svg"]:
            # Safely get element_id for map lookup
            raw_img_id_block = element.get("data-docgen-id")
            img_id_block_str: str | None = None
            if isinstance(raw_img_id_block, list):
                if raw_img_id_block:  # Check if list is not empty
                    img_id_block_str = str(raw_img_id_block[0])
            elif isinstance(raw_img_id_block, str):
                img_id_block_str = raw_img_id_block

            if img_id_block_str:
                element_map_data = self.elements_map.get(img_id_block_str, {})
                if element_map_data.get("geometry"):
                    target_for_image = None
                    if hasattr(
                        parent_docx_object, "add_paragraph"
                    ) and not isinstance(parent_docx_object, DocumentObject):
                        logger.debug(
                            "Image parent is a cell/similar, adding new "
                            "paragraph for image."
                        )
                        target_for_image = parent_docx_object
                    elif hasattr(parent_docx_object, "add_run"):
                        logger.debug("Image parent is a paragraph.")
                        target_for_image = parent_docx_object
                    else:
                        target_for_image = self.doc
                    self._add_element_as_image(
                        img_id_block_str, target_for_image
                    )
                else:
                    logger.debug(
                        f"Skip block img/svg {img_id_block_str}, no geometry."
                    )
            else:
                logger.error(
                    "Skip block <%s> no data-docgen-id: %.30s",
                    tag_name,
                    str(element),
                )
            return  # Consumed img/svg or skipped

        # Handle p, div, h* separately from other block/inline logic
        if tag_name in ["p", "div", "h1", "h2", "h3", "h4", "h5", "h6"]:
            current_paragraph = self._create_paragraph_for_block(
                parent_docx_object, indent_level_inches
            )

            if tag_name.startswith("h"):
                try:
                    level = int(tag_name[1])
                    if level == 1:
                        style_name = "Title"
                    elif level == 2:
                        style_name = "Subtitle"
                    else:
                        level = level - 2
                        style_name = f"Heading {level}"
                    current_paragraph.style = style_name
                except (ValueError, IndexError):
                    logger.warning(
                        "Could not apply heading style for %s", tag_name
                    )
                logger.debug(
                    "Applied style '%s' to paragraph for <%s>",
                    current_paragraph.style.name,  # type: ignore
                    tag_name,
                )

            # Get styles and classes for the block element itself
            element_styles, element_classes = self._parse_styles(element)
            new_active_tags = list(active_format_tags)
            # DIVs can provide formatting context via styles/classes
            if tag_name == "div":
                new_active_tags.append(element)
                # Still useful for style inheritance if not using classes
                # directly here

            for i, child in enumerate(element.children):
                self._process_inline_content(
                    child,
                    current_paragraph,
                    new_active_tags,
                    first=(i == 0),
                    is_dt_content=False,
                    crt_par_classes=element_classes,
                )

        elif tag_name in ["ul", "ol"]:
            for item in element.find_all("li", recursive=False):
                if not isinstance(item, Tag):
                    continue
                style = "ListBullet" if tag_name == "ul" else "ListNumber"
                logger.debug("Adding <li> item with style: %s", style)
                # Content of <li> processed into this new paragraph
                li_paragraph = self.doc.add_paragraph(style=style)
                if indent_level_inches:
                    li_paragraph.paragraph_format.left_indent = Inches(
                        indent_level_inches
                    )
                for i, child_content in enumerate(item.children):
                    self._process_inline_content(
                        child_content,
                        li_paragraph,
                        active_format_tags,
                        first=(i == 0),
                        is_dt_content=False,
                        crt_par_classes=crt_par_classes,
                    )

        elif tag_name == "dl":
            self._handle_dl_element(
                element,
                parent_docx_object,
                active_format_tags,
                indent_level_inches,
                crt_par_classes=crt_par_classes,
            )
            logger.debug("Finished processing <dl> element.")

        elif tag_name == "table":
            # Pass table classes from _parse_styles to _handle_table
            self._handle_table(element, parent_docx_object)

        elif tag_name == "hr":
            p = self.doc.add_paragraph()
            if indent_level_inches:
                p.paragraph_format.left_indent = Inches(indent_level_inches)
            pPr = p._p.get_or_add_pPr()
            pBdr, bottom = OxmlElement("w:pBdr"), OxmlElement("w:bottom")
            logger.debug("Adding <hr> as paragraph bottom border.")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # New explicit handling for known inline tags found at block
        # level
        elif tag_name in [
            "sup",
            "sub",
            "span",
            "font",
            "b",
            "i",
            "strong",
            "em",
            "u",
            "s",
            "a",
            "del",
            "strike",
            "br",
        ]:
            target_p_for_inline_tag = None
            is_first_run_in_target_p = True
            if hasattr(parent_docx_object, "_tc"):  # Parent is a cell
                if parent_docx_object.paragraphs:
                    target_p_for_inline_tag = parent_docx_object.paragraphs[-1]
                    is_first_run_in_target_p = not bool(
                        target_p_for_inline_tag.runs
                    )
                else:
                    target_p_for_inline_tag = parent_docx_object.add_paragraph()
            elif hasattr(
                parent_docx_object, "add_run"
            ):  # Parent is already a paragraph
                target_p_for_inline_tag = parent_docx_object
                is_first_run_in_target_p = not bool(
                    target_p_for_inline_tag.runs
                )
            else:  # Fallback
                target_p_for_inline_tag = self._create_paragraph_for_block(
                    parent_docx_object, indent_level_inches
                )

            self._process_inline_content(
                element,  # The inline Tag itself (e.g., <sup>, <br>)
                target_p_for_inline_tag,
                active_format_tags,  # Pass current active formatting context
                first=is_first_run_in_target_p,
                crt_par_classes=crt_par_classes,  # Pass through if available
            )

        else:
            # Default handling for unrecognized block tags or block-ish wrappers
            # (e.g. a div not handled above, or custom tags treated as blocks)
            new_active_tags = list(active_format_tags)
            if element.name.lower() in ["span", "font", "div"]:
                new_active_tags.append(element)

            # For children, decide if they are block (recursive call) or inline
            # (needs paragraph)
            created_para_for_inline = None
            for i, child in enumerate(element.children):
                if isinstance(child, Tag) and child.name.lower() in [
                    "p",
                    "div",
                    "h1",
                    "h2",
                    "h3",
                    "h4",
                    "h5",
                    "h6",
                    "ul",
                    "ol",
                    "table",
                    "img",
                    "svg",
                    "hr",
                    "dl",
                ]:
                    self._process_block_element(
                        child,
                        parent_docx_object,
                        new_active_tags,
                        indent_level_inches=indent_level_inches,
                        crt_par_classes=crt_par_classes,
                    )
                    created_para_for_inline = (
                        None  # Reset: next inline content needs new para
                    )
                elif isinstance(child, (NavigableString, Tag)):
                    # Inline or text found at block level, needs a paragraph
                    # context.
                    if created_para_for_inline is None:
                        # If parent_docx_object is a cell with existing
                        # paragraphs, and the current element (wrapper like
                        # <sup>) is inline, use the cell's last paragraph.
                        is_inline_wrapper = element.name.lower() in [
                            "sup",
                            "sub",
                            "span",
                            "font",
                            "b",
                            "i",
                            "strong",
                            "em",
                            "u",
                            "s",
                            "a",
                            "del",
                            "strike",
                        ]
                        if (
                            hasattr(parent_docx_object, "_tc")
                            and parent_docx_object.paragraphs
                            and is_inline_wrapper
                        ):
                            created_para_for_inline = (
                                parent_docx_object.paragraphs[-1]
                            )
                        else:
                            # Element is some other unknown block-ish wrapper,
                            # or parent is not a cell / has no paras / element
                            # is not inline wrapper.
                            created_para_for_inline = (
                                self._create_paragraph_for_block(
                                    parent_docx_object, indent_level_inches
                                )
                            )

                    self._process_inline_content(
                        child,
                        created_para_for_inline,
                        new_active_tags,
                        first=(i == 0)
                        and (
                            not bool(created_para_for_inline.runs)
                        ),  # Check if para is empty
                        is_dt_content=False,  # Assuming default context here
                        crt_par_classes=crt_par_classes,
                    )

    def _create_paragraph_for_block(
        self, parent_docx_object: Any, indent_level_inches: float | None
    ):
        """Helper to create a paragraph, typically for a new block element,
        applying indentation.
        """
        para = None
        if hasattr(parent_docx_object, "add_paragraph") and not isinstance(
            parent_docx_object, DocumentObject
        ):
            para = parent_docx_object.add_paragraph()  # Usually a Cell
            logger.debug(
                "Created paragraph in cell/block object for new block element."
            )
        else:
            # Document object, or parent is already a paragraph (less common
            # for new block)
            para = self.doc.add_paragraph()
            logger.debug(
                "Created paragraph in document root for new block element."
            )

        if indent_level_inches is not None and indent_level_inches > 0:
            para.paragraph_format.left_indent = Inches(indent_level_inches)
            logger.debug(
                "Applied indent of %s inches to new block paragraph.",
                indent_level_inches,
            )
        return para

    def _handle_dl_element(
        self,
        dl_element: Tag,
        parent_docx_object: Any,
        active_format_tags: list[Tag],
        current_indent_inches: float | None,
        crt_par_classes: list[str] | None = None,
    ):
        """Handles <dl> elements: styled <dt> and indented <dd> blocks."""
        logger.debug(
            "Processing <dl> element. Indent: %s", current_indent_inches
        )
        for child_node in dl_element.children:
            if not isinstance(child_node, Tag):
                continue

            child_tag_name = child_node.name.lower()

            if child_tag_name == "dt":
                logger.debug("  Processing <dt> in <dl>")
                p_dt = self._create_paragraph_for_block(
                    parent_docx_object, current_indent_inches
                )
                for i, dt_content_child in enumerate(child_node.children):
                    self._process_inline_content(
                        dt_content_child,
                        p_dt,
                        active_format_tags,
                        first=(i == 0),
                        is_dt_content=True,
                        crt_par_classes=crt_par_classes,
                    )

                # makes the dt paragraph have no spacing between it and the
                # following dd paragraph.
                p_dt.style = "No Spacing"

            elif child_tag_name == "dd":  # child_node is the <dd> Tag
                dd_children_indent_val = (current_indent_inches or 0) + 0.25
                logger.debug(
                    "  Processing <dd> in <dl>. New indent: %s",
                    dd_children_indent_val,
                )

                if not list(child_node.children):  # If DD is empty, skip
                    logger.debug("    <dd> is empty, skipping.")
                    continue

                # Create the primary paragraph for this <dd>'s inline content
                # stream
                current_dd_paragraph = self._create_paragraph_for_block(
                    parent_docx_object, dd_children_indent_val
                )
                is_first_run_in_dd_para = True

                for (
                    dd_content_item
                ) in child_node.children:  # Iterate children of <dd>
                    if isinstance(dd_content_item, (NavigableString, Tag)):
                        is_item_a_block_within_dd = False
                        if isinstance(dd_content_item, Tag):
                            item_tag_name = dd_content_item.name.lower()
                            # Define tags that should be treated as block-level
                            # when direct children of DD
                            if item_tag_name in [
                                "p",
                                "div",
                                "h1",
                                "h2",
                                "h3",
                                "h4",
                                "h5",
                                "h6",
                                "ul",
                                "ol",
                                "dl",
                                "table",
                                "hr",
                            ]:
                                logger.debug(
                                    "    <dd> child is block-level: <%s>",
                                    item_tag_name,
                                )
                                is_item_a_block_within_dd = True

                        if is_item_a_block_within_dd:
                            # Finalize current_dd_paragraph or remove if empty
                            if current_dd_paragraph:
                                if (
                                    not current_dd_paragraph.text.strip()
                                    and not current_dd_paragraph.runs
                                ):
                                    p_elem = current_dd_paragraph._element
                                    if p_elem.getparent() is not None:
                                        p_elem.getparent().remove(p_elem)
                                        logger.debug(
                                            "    Removed empty paragraph from "
                                            "<dd> before processing block "
                                            "child."
                                        )
                                # Set to None to signal a new one is needed if
                                # more inline content follows this block
                                current_dd_paragraph = None

                            # Process this block element. Its parent is the
                            # DL's parent. It will create its own paragraphs,
                            # using dd_children_indent_val.
                            self._process_block_element(
                                dd_content_item,  # The block Tag
                                parent_docx_object,
                                active_format_tags,
                                indent_level_inches=dd_children_indent_val,
                                crt_par_classes=crt_par_classes,
                            )
                            # After a block, the next inline item will need a
                            # new paragraph.
                            is_first_run_in_dd_para = (
                                True  # Reset for the potentially new paragraph
                            )
                            logger.debug(
                                "    Finished processing block child in <dd>. "
                                "Next inline needs new para."
                            )
                        else:
                            # This dd_content_item is inline (NavigableString
                            # or inline Tag like <sup>)
                            # It should go into the current_dd_paragraph.
                            logger.debug(
                                "    <dd> child is inline: %s",
                                str(dd_content_item)[:50].replace("\n", " "),
                            )
                            if (
                                current_dd_paragraph is None
                            ):  # Means a block was just processed
                                current_dd_paragraph = (
                                    self._create_paragraph_for_block(
                                        parent_docx_object,
                                        dd_children_indent_val,
                                    )
                                )
                                is_first_run_in_dd_para = True

                            self._process_inline_content(
                                # The NavigableString or inline Tag
                                dd_content_item,
                                current_dd_paragraph,
                                active_format_tags,
                                first=is_first_run_in_dd_para,
                                crt_par_classes=crt_par_classes,
                            )
                            # Update is_first_run_in_dd_para based on whether
                            # content was added
                            if (
                                isinstance(dd_content_item, NavigableString)
                                and str(dd_content_item).strip()
                            ):
                                is_first_run_in_dd_para = False
                            elif isinstance(dd_content_item, Tag):
                                # Consider content-ful tags or explicit line
                                # breaks as "content added"
                                if dd_content_item.get_text(
                                    strip=True
                                ) or dd_content_item.name.lower() in [
                                    "br",
                                    "img",
                                    "svg",
                                ]:
                                    is_first_run_in_dd_para = False
                                    logger.debug(
                                        "    Inline content added to <dd> "
                                        "paragraph, is_first_run_in_dd_para "
                                        "set to False."
                                    )
                    # else: ignore other PageElement types like Comment, etc.

                # After processing all children of <dd>, if the last
                # current_dd_paragraph is empty, remove it.
                if (
                    current_dd_paragraph
                    and not current_dd_paragraph.text.strip()
                    and not current_dd_paragraph.runs
                ):
                    p_elem = current_dd_paragraph._element
                    if (
                        p_elem.getparent() is not None
                    ):  # Ensure it's still part of the tree
                        p_elem.getparent().remove(p_elem)
                        logger.debug(
                            "  Removed last empty paragraph from <dd> after "
                            "processing all its children."
                        )
        logger.debug("Finished processing <dl> element.")

    def _get_attribute_as_int(
        self, element: Tag, attr_name: str, default_val: int = 1
    ) -> int:
        """Helper to safely get an attribute value as an integer."""
        attr_val = element.get(attr_name)
        if attr_val is None:
            return default_val

        val_to_convert = attr_val
        if isinstance(attr_val, list):
            if not attr_val:  # Empty list
                return default_val
            val_to_convert = str(attr_val[0])  # Take first item if list
        else:
            val_to_convert = str(attr_val)

        try:
            return int(val_to_convert)
        except ValueError:
            logger.warning(
                f"Invalid value for attribute {attr_name}: {val_to_convert}. "
                f"Using default {default_val}."
            )
            return default_val
